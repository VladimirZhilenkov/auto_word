"""
Order Edit Dialog - Edit existing orders from the journal.

Shows template preview with highlighted variables and allows
editing listeners, order parameters, and regenerating the document.
"""

import html as html_module
import re
import subprocess
import sys
from datetime import datetime, date
from pathlib import Path
from typing import List, Optional, Dict, Any

from PyQt5.QtCore import Qt, QDate
from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QFormLayout, QSplitter,
    QPushButton, QLabel, QGroupBox, QMessageBox, QFrame,
    QComboBox, QListWidget, QListWidgetItem, QTextBrowser, QTextEdit,
    QLineEdit, QDateEdit, QSpinBox, QTabWidget, QCheckBox,
    QAbstractItemView, QApplication, QWidget, QScrollArea,
    QSizePolicy, QFileDialog, QInputDialog
)
from PyQt5.QtGui import QFont, QColor

from ...database.connection import DatabaseSession
from ...database.models import Listener, Program, ProgramListener
from ...services.document_generator import DocumentGenerator
from ...services.order_journal_service import OrderJournalService, ORDER_TYPE_LABELS


# ---------------------------------------------------------------------------
# Template text extraction with variable highlighting
# ---------------------------------------------------------------------------

def _highlight_variables(text: str) -> str:
    """Highlight Jinja2 variables and tags in text with HTML spans."""
    text = html_module.escape(text)

    # {{ variable }} → blue
    text = re.sub(
        r'\{\{\s*(.+?)\s*\}\}',
        r'<span style="background-color:#BBDEFB; color:#0D47A1; '
        r'padding:1px 4px; border-radius:3px; font-weight:bold;">'
        r'{{ \1 }}</span>',
        text,
    )

    # {% tag %} → green
    text = re.sub(
        r'\{%\s*(.+?)\s*%\}',
        r'<span style="background-color:#C8E6C9; color:#1B5E20; '
        r'padding:1px 4px; border-radius:3px; font-weight:bold;">'
        r'{% \1 %}</span>',
        text,
    )

    # {# comment #} → yellow
    text = re.sub(
        r'\{#\s*(.+?)\s*#\}',
        r'<span style="background-color:#FFF9C4; color:#F57F17; '
        r'padding:1px 4px; border-radius:3px;">'
        r'{# \1 #}</span>',
        text,
    )

    return text


def extract_template_html(template_path: str) -> str:
    """
    Extract text from a .docx template and return HTML with
    highlighted Jinja2 variables/tags.
    """
    try:
        from docx import Document
    except ImportError:
        return "<p style='color:red;'>python-docx не установлен</p>"

    path = Path(template_path)
    if not path.exists():
        return f"<p style='color:red;'>Файл не найден: {html_module.escape(str(path))}</p>"

    try:
        doc = Document(str(path))
    except Exception as exc:
        return f"<p style='color:red;'>Ошибка открытия: {html_module.escape(str(exc))}</p>"

    parts: List[str] = []
    parts.append(
        "<style>"
        "body { font-family: 'Times New Roman', serif; font-size: 13px; }"
        "table { border-collapse: collapse; width: 100%; margin: 8px 0; }"
        "td, th { border: 1px solid #999; padding: 4px 6px; }"
        "p { margin: 2px 0; }"
        ".empty { color: #999; font-style: italic; }"
        "</style>"
    )

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            highlighted = _highlight_variables(text)
            # Detect alignment from paragraph
            alignment = ""
            if para.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = " style='text-align:center;'"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = " style='text-align:right;'"
            # If paragraph is bold
            is_bold = any(run.bold for run in para.runs if run.bold)
            if is_bold:
                highlighted = f"<b>{highlighted}</b>"
            parts.append(f"<p{alignment}>{highlighted}</p>")

    # Tables
    for table in doc.tables:
        parts.append("<table>")
        for row_idx, row in enumerate(table.rows):
            parts.append("<tr>")
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    highlighted = _highlight_variables(cell_text)
                else:
                    highlighted = '<span class="empty">—</span>'
                tag = "th" if row_idx == 0 else "td"
                parts.append(f"<{tag}>{highlighted}</{tag}>")
            parts.append("</tr>")
        parts.append("</table>")

    if len(parts) <= 1:  # only style tag
        return "<p class='empty'>Шаблон пуст или не содержит текста</p>"

    return "\n".join(parts)


def extract_template_variables(template_path: str) -> List[str]:
    """Return sorted list of unique Jinja2 variable names from a template."""
    try:
        from docx import Document
        doc = Document(str(template_path))
    except Exception:
        return []

    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"

    variables = set()
    for m in re.finditer(r'\{\{\s*([\w.]+)\s*\}\}', full_text):
        variables.add(m.group(1))

    return sorted(variables)


def extract_template_plain_text(template_path: str) -> str:
    """Extract plain text from a .docx template for editing."""
    try:
        from docx import Document
        doc = Document(str(template_path))
    except Exception as exc:
        return f"(Ошибка открытия: {exc})"

    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            parts.append(text)
        else:
            parts.append("")  # keep empty lines

    for table in doc.tables:
        parts.append("\n--- ТАБЛИЦА ---")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
        parts.append("--- /ТАБЛИЦА ---\n")

    return "\n".join(parts)


def save_text_to_template(template_path: str, new_text: str) -> None:
    """
    Save edited text back into a .docx template, preserving formatting.

    Strategy: map edited paragraphs to original paragraphs by index.
    For table sections, map edited rows back to table cells.
    Creates backup before saving.
    """
    import shutil
    from docx import Document

    path = Path(template_path)
    # Create backup
    backup_path = path.with_suffix('.docx.bak')
    shutil.copy2(str(path), str(backup_path))

    doc = Document(str(path))

    lines = new_text.split('\n')

    # Separate paragraph lines and table sections
    para_lines: List[str] = []
    table_sections: List[List[str]] = []
    current_table: Optional[List[str]] = None

    for line in lines:
        if line.strip() == '--- ТАБЛИЦА ---':
            current_table = []
        elif line.strip() == '--- /ТАБЛИЦА ---':
            if current_table is not None:
                table_sections.append(current_table)
            current_table = None
        elif current_table is not None:
            current_table.append(line)
        else:
            para_lines.append(line)

    # Update paragraphs
    for i, para in enumerate(doc.paragraphs):
        if i < len(para_lines):
            new_line = para_lines[i]
            if para.text != new_line:
                _replace_paragraph_text(para, new_line)

    # Update tables
    for t_idx, table in enumerate(doc.tables):
        if t_idx < len(table_sections):
            table_lines = table_sections[t_idx]
            for r_idx, row in enumerate(table.rows):
                if r_idx < len(table_lines):
                    cells_text = table_lines[r_idx].split(' | ')
                    for c_idx, cell in enumerate(row.cells):
                        if c_idx < len(cells_text):
                            new_cell_text = cells_text[c_idx].strip()
                            if cell.text.strip() != new_cell_text:
                                _replace_cell_text(cell, new_cell_text)

    doc.save(str(path))


def _replace_paragraph_text(para, new_text: str):
    """Replace paragraph text while preserving run formatting."""
    if not para.runs:
        # No runs - just set text
        para.text = new_text
        return

    # Strategy: put all text in first run, clear others
    if len(para.runs) == 1:
        para.runs[0].text = new_text
    else:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def _replace_cell_text(cell, new_text: str):
    """Replace cell text while preserving formatting."""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        _replace_paragraph_text(para, new_text)
        # Clear extra paragraphs
        for p in cell.paragraphs[1:]:
            for run in p.runs:
                run.text = ""
    else:
        cell.text = new_text


# ---------------------------------------------------------------------------
# Complete variable reference data
# ---------------------------------------------------------------------------

VARIABLE_REFERENCE = {
    "📅 Данные приказа": [
        {"var": "order_number", "desc": "Номер приказа", "example": "111"},
        {"var": "order_day", "desc": "День приказа (число)", "example": "12"},
        {"var": "order_month", "desc": "Месяц приказа (текстом)", "example": "ноября"},
        {"var": "order_year", "desc": "Год приказа", "example": "2025"},
        {"var": "order_type_label", "desc": "Полное название типа приказа", "example": "О зачислении слушателей"},
    ],
    "📜 Данные договора/контракта": [
        {"var": "contract_type", "desc": "Тип договора/контракта", "example": "государственным контрактом"},
        {"var": "contract_number", "desc": "Номер договора", "example": "б/н"},
        {"var": "contract_date", "desc": "Дата договора", "example": "от 28 мая 2025 года"},
    ],
    "🎓 Программа обучения": [
        {"var": "program_name", "desc": "Полное название программы", "example": "Организация работы администраторов районных судов"},
        {"var": "program_short_name", "desc": "Краткое название программы", "example": "Админ. районных судов"},
        {"var": "program_type", "desc": "Вид программы", "example": "повышения квалификации"},
        {"var": "stream_name", "desc": "Название потока", "example": "государственных гражданских служащих"},
        {"var": "hours", "desc": "Количество часов", "example": "16"},
        {"var": "education_form", "desc": "Форма обучения", "example": "очная"},
        {"var": "education_form_genitive", "desc": "Форма обучения (родительный падеж)", "example": "очной"},
        {"var": "education_format", "desc": "Формат обучения", "example": "с применением электронного обучения"},
        {"var": "start_date", "desc": "Дата начала обучения", "example": "20 января 2026 г."},
        {"var": "end_date", "desc": "Дата окончания обучения", "example": "25 января 2026 г."},
        {"var": "training_period", "desc": "Период обучения (текст)", "example": "20.01.2026 - 25.01.2026"},
        {"var": "training_duration", "desc": "Продолжительность обучения", "example": "6 дней"},
        {"var": "listeners_count", "desc": "Количество слушателей", "example": "15"},
    ],
    "👥 Слушатель (в таблице, с циклом)": [
        {"var": "loop.index", "desc": "Порядковый номер (1, 2, 3...)", "example": "1"},
        {"var": "listener.full_name", "desc": "ФИО полностью", "example": "Иванов Иван Иванович"},
        {"var": "listener.last_name", "desc": "Фамилия", "example": "Иванов"},
        {"var": "listener.first_name", "desc": "Имя", "example": "Иван"},
        {"var": "listener.middle_name", "desc": "Отчество", "example": "Иванович"},
        {"var": "listener.position", "desc": "Должность", "example": "Администратор суда"},
        {"var": "listener.court_name", "desc": "Наименование суда (место работы)", "example": "Ленинский районный суд г. Иркутска"},
        {"var": "listener.region", "desc": "Субъект РФ", "example": "Иркутская область"},
        {"var": "listener.order_number", "desc": "Порядковый номер слушателя", "example": "1"},
        {"var": "listener.contract_number", "desc": "Номер договора слушателя", "example": "42"},
        {"var": "listener.contract_date", "desc": "Дата договора слушателя", "example": "15.01.2026"},
        {"var": "listener.notes", "desc": "Примечание слушателя", "example": ""},
    ],
    "👤 Слушатель (индивидуальный, именительный)": [
        {"var": "full_name", "desc": "ФИО (именительный)", "example": "Иванов Иван Иванович"},
        {"var": "last_name", "desc": "Фамилия (именительный)", "example": "Иванов"},
        {"var": "first_name", "desc": "Имя (именительный)", "example": "Иван"},
        {"var": "middle_name", "desc": "Отчество (именительный)", "example": "Иванович"},
        {"var": "initials", "desc": "Фамилия И.О.", "example": "Иванов И.И."},
        {"var": "initials_before", "desc": "И.О. Фамилия", "example": "И.И. Иванов"},
        {"var": "position", "desc": "Должность", "example": "Администратор суда"},
        {"var": "workplace", "desc": "Место работы / суд", "example": "Ленинский районный суд"},
        {"var": "court_name", "desc": "Наименование суда (алиас workplace)", "example": "Ленинский районный суд"},
        {"var": "region", "desc": "Субъект РФ", "example": "Иркутская область"},
    ],
    "📝 Склонение ФИО (падежи)": [
        {"var": "full_name_genitive", "desc": "ФИО (родительный — кого?)", "example": "Иванова Ивана Ивановича"},
        {"var": "full_name_dative", "desc": "ФИО (дательный — кому?)", "example": "Иванову Ивану Ивановичу"},
        {"var": "full_name_accusative", "desc": "ФИО (винительный — кого?)", "example": "Иванова Ивана Ивановича"},
        {"var": "full_name_instrumental", "desc": "ФИО (творительный — кем?)", "example": "Ивановым Иваном Ивановичем"},
        {"var": "full_name_prepositional", "desc": "ФИО (предложный — о ком?)", "example": "Иванове Иване Ивановиче"},
        {"var": "last_name_genitive", "desc": "Фамилия (родительный)", "example": "Иванова"},
        {"var": "last_name_dative", "desc": "Фамилия (дательный)", "example": "Иванову"},
        {"var": "first_name_genitive", "desc": "Имя (родительный)", "example": "Ивана"},
        {"var": "first_name_dative", "desc": "Имя (дательный)", "example": "Ивану"},
        {"var": "middle_name_genitive", "desc": "Отчество (родительный)", "example": "Ивановича"},
        {"var": "middle_name_dative", "desc": "Отчество (дательный)", "example": "Ивановичу"},
        {"var": "initials_genitive", "desc": "Фамилия И.О. (родительный)", "example": "Иванова И.И."},
        {"var": "initials_dative", "desc": "Фамилия И.О. (дательный)", "example": "Иванову И.И."},
    ],
    "📞 Контактные данные": [
        {"var": "birth_date", "desc": "Дата рождения", "example": "15.03.1985"},
        {"var": "mobile_phone", "desc": "Мобильный телефон", "example": "+7 (999) 123-45-67"},
        {"var": "work_phone", "desc": "Рабочий телефон", "example": "+7 (495) 123-45-67"},
        {"var": "email", "desc": "Электронная почта", "example": "ivanov@example.com"},
    ],
    "🪪 Паспортные данные": [
        {"var": "passport_series_number", "desc": "Серия и номер паспорта", "example": "1234 567890"},
        {"var": "passport_issue_date", "desc": "Дата выдачи паспорта", "example": "25.12.2010"},
        {"var": "passport_issued_by", "desc": "Кем выдан паспорт", "example": "ОВД Ленинского района"},
        {"var": "passport_department_code", "desc": "Код подразделения", "example": "380-015"},
    ],
    "🏠 Адреса и идентификаторы": [
        {"var": "registration_address", "desc": "Адрес регистрации с индексом", "example": "664025, г. Иркутск, ул. Ленина, д. 10"},
        {"var": "actual_address", "desc": "Фактический адрес проживания", "example": "664025, г. Иркутск, ул. Ленина, д. 10"},
        {"var": "snils", "desc": "СНИЛС", "example": "123-456-789 00"},
        {"var": "inn", "desc": "ИНН", "example": "123456789012"},
        {"var": "personal_data_consent", "desc": "Согласие на обработку ПД", "example": "Да"},
    ],
    "👔 Исполнитель": [
        {"var": "executor_name", "desc": "ФИО исполнителя", "example": "Петрова А.В."},
        {"var": "executor_position", "desc": "Должность исполнителя", "example": "Диспетчер ФПК"},
        {"var": "executor_full", "desc": "Должность + ФИО исполнителя", "example": "Диспетчер ФПК Петрова А.В."},
    ],
    "📆 Служебные": [
        {"var": "current_date", "desc": "Текущая дата генерации", "example": "03.03.2026"},
        {"var": "current_year", "desc": "Текущий год", "example": "2026"},
        {"var": "current_month", "desc": "Текущий месяц (число)", "example": "03"},
        {"var": "current_day", "desc": "Текущий день (число)", "example": "03"},
        {"var": "generation_datetime", "desc": "Дата и время генерации", "example": "03.03.2026 14:30:00"},
    ],
}


# ---------------------------------------------------------------------------
# Order Edit Dialog
# ---------------------------------------------------------------------------

class OrderEditDialog(QDialog):
    """
    Dialog for editing an existing order from the journal.

    Features:
      - Template preview with highlighted variables
      - Listener list with add/remove capability
      - Order parameter editing
      - Document regeneration
    """

    def __init__(
        self,
        parent=None,
        entry: Optional[Dict[str, Any]] = None,
    ):
        super().__init__(parent)
        self.entry = entry or {}
        self.generator = DocumentGenerator()
        self.journal_service = OrderJournalService()

        self._all_programs: List[Dict[str, Any]] = []
        self._program_listeners: List[Dict[str, Any]] = []
        self._extra_listeners: List[Dict[str, Any]] = []  # listeners added manually

        self._setup_ui()
        self._load_data()
        self._update_template_preview()

    # ------------------------------------------------------------------
    # UI setup
    # ------------------------------------------------------------------

    def _setup_ui(self):
        self.setWindowTitle("Редактирование приказа")
        self.setMinimumSize(900, 600)
        self.setModal(True)

        # Auto-size to ~85% of screen
        screen = QApplication.primaryScreen()
        if screen:
            geom = screen.availableGeometry()
            w = int(geom.width() * 0.85)
            h = int(geom.height() * 0.85)
            self.resize(w, h)
            # Center on screen
            x = geom.x() + (geom.width() - w) // 2
            y = geom.y() + (geom.height() - h) // 2
            self.move(x, y)
        else:
            self.resize(1400, 900)

        root_layout = QVBoxLayout(self)

        # Main splitter: left = template preview, right = editing panel
        splitter = QSplitter(Qt.Horizontal)

        # --- LEFT: Template Panel with Tabs ---
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)

        # Template selector (shared across tabs)
        tpl_select_layout = QHBoxLayout()
        tpl_select_layout.addWidget(QLabel("Шаблон:"))
        self.combo_template = QComboBox()
        self.combo_template.setMinimumWidth(250)
        self.combo_template.currentIndexChanged.connect(self._on_template_changed)
        tpl_select_layout.addWidget(self.combo_template, stretch=1)

        self.btn_refresh_templates = QPushButton("🔄")
        self.btn_refresh_templates.setMaximumWidth(36)
        self.btn_refresh_templates.setToolTip("Обновить список шаблонов")
        self.btn_refresh_templates.clicked.connect(self._load_templates)
        tpl_select_layout.addWidget(self.btn_refresh_templates)

        self.btn_open_templates_folder = QPushButton("📂 Папка шаблонов")
        self.btn_open_templates_folder.clicked.connect(self._open_templates_folder)
        tpl_select_layout.addWidget(self.btn_open_templates_folder)

        self.btn_fix_templates = QPushButton("🔧 Исправить все шаблоны")
        self.btn_fix_templates.setToolTip(
            "Добавить цикл {% for listener in listeners %} во все шаблоны,\n"
            "где он отсутствует (необходим для корректной генерации)"
        )
        self.btn_fix_templates.clicked.connect(self._fix_all_templates)
        tpl_select_layout.addWidget(self.btn_fix_templates)

        left_layout.addLayout(tpl_select_layout)

        # Tabs: Preview / Editor / Reference
        self.left_tabs = QTabWidget()

        # -- Tab 1: Preview --
        preview_tab = QWidget()
        preview_layout = QVBoxLayout(preview_tab)
        preview_layout.setContentsMargins(2, 2, 2, 2)

        self.template_browser = QTextBrowser()
        self.template_browser.setOpenExternalLinks(False)
        self.template_browser.setFont(QFont("Times New Roman", 11))
        preview_layout.addWidget(self.template_browser)

        # Variables legend
        legend = QLabel(
            '<span style="background:#BBDEFB;color:#0D47A1;padding:2px 6px;border-radius:3px;">'
            '{{ переменная }}</span> &nbsp; '
            '<span style="background:#C8E6C9;color:#1B5E20;padding:2px 6px;border-radius:3px;">'
            '{% управление %}</span>'
        )
        legend.setTextFormat(Qt.RichText)
        legend.setStyleSheet("padding: 4px;")
        preview_layout.addWidget(legend)

        self.label_variables = QLabel("")
        self.label_variables.setWordWrap(True)
        self.label_variables.setStyleSheet("color: #555; font-size: 11px; padding: 4px;")
        preview_layout.addWidget(self.label_variables)

        self.left_tabs.addTab(preview_tab, "📄 Предпросмотр")

        # -- Tab 2: Template Editor --
        editor_tab = QWidget()
        editor_layout = QVBoxLayout(editor_tab)
        editor_layout.setContentsMargins(2, 2, 2, 2)

        # Quick insert bar
        insert_group = QGroupBox("Быстрая вставка переменной")
        insert_layout = QVBoxLayout(insert_group)

        # Category selector
        cat_layout = QHBoxLayout()
        cat_layout.addWidget(QLabel("Категория:"))
        self.combo_var_category = QComboBox()
        self.combo_var_category.currentIndexChanged.connect(self._on_var_category_changed)
        cat_layout.addWidget(self.combo_var_category, stretch=1)
        insert_layout.addLayout(cat_layout)

        # Variable selector
        var_layout = QHBoxLayout()
        var_layout.addWidget(QLabel("Переменная:"))
        self.combo_var_name = QComboBox()
        self.combo_var_name.setMinimumWidth(200)
        self.combo_var_name.currentIndexChanged.connect(self._on_var_selected)
        var_layout.addWidget(self.combo_var_name, stretch=1)
        insert_layout.addLayout(var_layout)

        # Variable description
        self.label_var_desc = QLabel("")
        self.label_var_desc.setWordWrap(True)
        self.label_var_desc.setStyleSheet(
            "color: #555; font-style: italic; padding: 2px 4px; font-size: 11px;"
        )
        insert_layout.addWidget(self.label_var_desc)

        # Insert buttons
        insert_btn_layout = QHBoxLayout()
        self.btn_copy_var = QPushButton("📋 Копировать")
        self.btn_copy_var.setToolTip("Скопировать переменную в буфер обмена")
        self.btn_copy_var.clicked.connect(self._copy_variable)
        insert_btn_layout.addWidget(self.btn_copy_var)

        self.btn_insert_var = QPushButton("⬇️ Вставить в текст")
        self.btn_insert_var.setToolTip("Вставить переменную в позицию курсора в редакторе ниже")
        self.btn_insert_var.clicked.connect(self._insert_variable)
        insert_btn_layout.addWidget(self.btn_insert_var)
        insert_layout.addLayout(insert_btn_layout)

        editor_layout.addWidget(insert_group)

        # Template text editor
        edit_text_group = QGroupBox("Редактор текста шаблона")
        edit_text_layout = QVBoxLayout(edit_text_group)

        self.template_editor = QTextEdit()
        self.template_editor.setFont(QFont("Courier New", 11))
        self.template_editor.setPlaceholderText(
            "Текст шаблона загрузится автоматически при выборе шаблона.\n"
            "Вы можете редактировать текст и переменные здесь."
        )
        edit_text_layout.addWidget(self.template_editor)

        # Editor action buttons
        editor_btn_layout = QHBoxLayout()

        self.btn_add_variable = QPushButton("➕ Добавить переменную")
        self.btn_add_variable.setToolTip("Добавить новую переменную в позицию курсора")
        self.btn_add_variable.clicked.connect(self._add_variable_at_cursor)
        editor_btn_layout.addWidget(self.btn_add_variable)

        self.btn_replace_variable = QPushButton("🔄 Заменить переменную")
        self.btn_replace_variable.setToolTip("Найти и заменить переменную в тексте шаблона")
        self.btn_replace_variable.clicked.connect(self._replace_variable)
        editor_btn_layout.addWidget(self.btn_replace_variable)

        self.btn_save_template = QPushButton("💾 Сохранить шаблон")
        self.btn_save_template.setStyleSheet(
            "font-weight: bold; background-color: #4CAF50; color: white;"
        )
        self.btn_save_template.clicked.connect(self._save_template)
        editor_btn_layout.addWidget(self.btn_save_template)

        edit_text_layout.addLayout(editor_btn_layout)
        editor_layout.addWidget(edit_text_group)

        self.left_tabs.addTab(editor_tab, "✏️ Редактор шаблона")

        # -- Tab 3: Variable Reference --
        reference_tab = QWidget()
        reference_layout = QVBoxLayout(reference_tab)
        reference_layout.setContentsMargins(2, 2, 2, 2)

        self.reference_browser = QTextBrowser()
        self.reference_browser.setOpenExternalLinks(False)
        self.reference_browser.setFont(QFont("Segoe UI", 10))
        self.reference_browser.setHtml(self._build_variable_reference_html())
        reference_layout.addWidget(self.reference_browser)

        self.left_tabs.addTab(reference_tab, "📚 Справка переменных")

        left_layout.addWidget(self.left_tabs)

        # Populate variable categories
        self._populate_var_categories()

        splitter.addWidget(left_widget)

        # --- RIGHT: Editing Panel ---
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # Scrollable area for all controls
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)

        # -- Order details --
        details_group = QGroupBox("Реквизиты приказа")
        details_form = QFormLayout(details_group)

        self.combo_order_type = QComboBox()
        for key, name in ORDER_TYPE_LABELS.items():
            self.combo_order_type.addItem(name, key)
        self.combo_order_type.currentIndexChanged.connect(self._on_type_changed)
        details_form.addRow("Тип:", self.combo_order_type)

        number_layout = QVBoxLayout()
        self.edit_order_number = QLineEdit()
        self.edit_order_number.setPlaceholderText("Номер приказа")
        number_layout.addWidget(self.edit_order_number)
        self.label_number_hint = QLabel("")
        self.label_number_hint.setStyleSheet("color: gray; font-size: 10px;")
        number_layout.addWidget(self.label_number_hint)
        details_form.addRow("Номер приказа:", number_layout)

        self.edit_order_date = QDateEdit()
        self.edit_order_date.setCalendarPopup(True)
        self.edit_order_date.setDisplayFormat("dd.MM.yyyy")
        details_form.addRow("Дата приказа:", self.edit_order_date)

        self.edit_title = QLineEdit()
        details_form.addRow("Наименование:", self.edit_title)

        # Executor position (combo with predefined values)
        self.combo_executor_position = QComboBox()
        self.combo_executor_position.setEditable(True)
        self.combo_executor_position.addItem(
            "Диспетчер факультета повышения квалификации и переподготовки судей, "
            "государственных гражданских служащих судов и Судебного департамента "
            "(ФПК судей и госслужащих судов)"
        )
        self.combo_executor_position.addItem(
            "Специалист по учебной работе I категории факультета повышения квалификации и переподготовки судей, "
            "государственных гражданских служащих судов и Судебного департамента "
            "(ФПК судей и госслужащих судов)"
        )
        details_form.addRow("Должность исполнителя:", self.combo_executor_position)

        # Executor name
        self.edit_executor = QLineEdit()
        self.edit_executor.setPlaceholderText("ФИО исполнителя")
        details_form.addRow("Имя исполнителя:", self.edit_executor)

        scroll_layout.addWidget(details_group)

        # -- Program details --
        program_group = QGroupBox("Программа обучения")
        program_form = QFormLayout(program_group)

        self.combo_program = QComboBox()
        self.combo_program.setMinimumWidth(300)
        self.combo_program.currentIndexChanged.connect(self._on_program_changed)
        program_form.addRow("Программа:", self.combo_program)

        # Program type
        self.combo_program_type = QComboBox()
        self.combo_program_type.addItem("повышения квалификации", "повышения квалификации")
        self.combo_program_type.addItem("профессиональной переподготовки", "профессиональной переподготовки")
        program_form.addRow("Вид программы:", self.combo_program_type)

        # Education form
        self.combo_education_form = QComboBox()
        self.combo_education_form.addItem("очная", "очная")
        self.combo_education_form.addItem("заочная", "заочная")
        self.combo_education_form.addItem("очно-заочная", "очно-заочная")
        program_form.addRow("Форма обучения:", self.combo_education_form)

        # Education format
        self.combo_education_format = QComboBox()
        self.combo_education_format.setEditable(True)
        self.combo_education_format.addItem("", "")
        self.combo_education_format.addItem(
            "с применением электронного обучения",
            "с применением электронного обучения",
        )
        self.combo_education_format.addItem(
            "с применением дистанционных образовательных технологий",
            "с применением дистанционных образовательных технологий",
        )
        self.combo_education_format.addItem(
            "с применением электронного обучения и дистанционных образовательных технологий",
            "с применением электронного обучения и дистанционных образовательных технологий",
        )
        program_form.addRow("Формат обучения:", self.combo_education_format)

        # Dates & hours
        dates_layout = QHBoxLayout()
        self.edit_start_date = QDateEdit()
        self.edit_start_date.setCalendarPopup(True)
        self.edit_start_date.setDisplayFormat("dd.MM.yyyy")
        self.edit_start_date.setDate(QDate.currentDate())
        dates_layout.addWidget(QLabel("Начало:"))
        dates_layout.addWidget(self.edit_start_date)

        self.edit_end_date = QDateEdit()
        self.edit_end_date.setCalendarPopup(True)
        self.edit_end_date.setDisplayFormat("dd.MM.yyyy")
        self.edit_end_date.setDate(QDate.currentDate())
        dates_layout.addWidget(QLabel("Окончание:"))
        dates_layout.addWidget(self.edit_end_date)

        self.spin_hours = QSpinBox()
        self.spin_hours.setRange(1, 2000)
        self.spin_hours.setValue(16)
        dates_layout.addWidget(QLabel("Часов:"))
        dates_layout.addWidget(self.spin_hours)
        program_form.addRow("Период:", dates_layout)

        scroll_layout.addWidget(program_group)

        # -- Listeners --
        listeners_group = QGroupBox("Слушатели")
        listeners_layout = QVBoxLayout(listeners_group)

        self.list_listeners = QListWidget()
        self.list_listeners.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.list_listeners.setMinimumHeight(180)
        listeners_layout.addWidget(self.list_listeners)

        # Listener action buttons
        listener_btn_layout = QHBoxLayout()

        self.btn_select_all = QPushButton("✅ Выбрать всех")
        self.btn_select_all.clicked.connect(self._select_all_listeners)
        listener_btn_layout.addWidget(self.btn_select_all)

        self.btn_deselect_all = QPushButton("❌ Снять выбор")
        self.btn_deselect_all.clicked.connect(self._deselect_all_listeners)
        listener_btn_layout.addWidget(self.btn_deselect_all)

        self.btn_add_listener = QPushButton("➕ Добавить слушателя")
        self.btn_add_listener.clicked.connect(self._add_listener)
        listener_btn_layout.addWidget(self.btn_add_listener)

        self.btn_remove_listener = QPushButton("➖ Убрать слушателя")
        self.btn_remove_listener.clicked.connect(self._remove_listener)
        listener_btn_layout.addWidget(self.btn_remove_listener)

        listeners_layout.addLayout(listener_btn_layout)

        self.label_listeners_count = QLabel("Выбрано: 0")
        self.label_listeners_count.setStyleSheet("font-weight: bold; padding: 2px;")
        listeners_layout.addWidget(self.label_listeners_count)

        # Connect item changed signal for auto-count
        self.list_listeners.itemChanged.connect(self._update_listener_count)

        scroll_layout.addWidget(listeners_group)

        # -- Notes --
        notes_group = QGroupBox("Примечание")
        notes_layout = QVBoxLayout(notes_group)
        self.edit_notes = QLineEdit()
        notes_layout.addWidget(self.edit_notes)
        scroll_layout.addWidget(notes_group)

        # -- Options --
        options_group = QGroupBox("Параметры")
        options_layout = QVBoxLayout(options_group)

        self.check_open_folder = QCheckBox("Открыть папку после генерации")
        self.check_open_folder.setChecked(True)
        options_layout.addWidget(self.check_open_folder)

        scroll_layout.addWidget(options_group)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        right_layout.addWidget(scroll)

        splitter.addWidget(right_widget)

        # Splitter proportions: 45% template, 55% editing
        splitter.setStretchFactor(0, 45)
        splitter.setStretchFactor(1, 55)

        root_layout.addWidget(splitter)

        # -- Bottom action buttons --
        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setFrameShadow(QFrame.Sunken)
        root_layout.addWidget(sep)

        bottom_layout = QHBoxLayout()

        self.btn_open_document = QPushButton("📂 Открыть документ")
        self.btn_open_document.clicked.connect(self._open_document)
        bottom_layout.addWidget(self.btn_open_document)

        bottom_layout.addStretch()

        self.btn_save_journal = QPushButton("💾 Сохранить запись журнала")
        self.btn_save_journal.clicked.connect(self._save_journal_entry)
        bottom_layout.addWidget(self.btn_save_journal)

        self.btn_regenerate = QPushButton("🔄 Перегенерировать документ")
        self.btn_regenerate.setStyleSheet(
            "font-weight: bold; padding: 6px 16px; background-color: #4CAF50; color: white;"
        )
        self.btn_regenerate.clicked.connect(self._regenerate_document)
        bottom_layout.addWidget(self.btn_regenerate)

        self.btn_close = QPushButton("Закрыть")
        self.btn_close.clicked.connect(self.reject)
        bottom_layout.addWidget(self.btn_close)

        root_layout.addLayout(bottom_layout)

    # ------------------------------------------------------------------
    # Data loading
    # ------------------------------------------------------------------

    def _load_data(self):
        """Load all data: programs, templates, and fill form from entry."""
        self._load_programs()
        self._load_templates()
        self._fill_from_entry()

    def _load_programs(self):
        """Load programs into combobox."""
        self.combo_program.blockSignals(True)
        self.combo_program.clear()
        self.combo_program.addItem("-- Без программы --", None)
        try:
            with DatabaseSession() as session:
                programs = session.query(Program).order_by(Program.program_name).all()
                for p in programs:
                    label = p.program_short_name or p.program_name
                    if len(label) > 80:
                        label = label[:80] + "..."
                    self.combo_program.addItem(label, p.id)
                    self._all_programs.append({
                        'id': p.id,
                        'program_name': p.program_name,
                        'program_short_name': p.program_short_name,
                        'training_period': p.training_period,
                        'training_duration': p.training_duration,
                        'program_volume': p.program_volume,
                        'education_form': p.education_form,
                        'education_format': p.education_format,
                    })
        except Exception as exc:
            QMessageBox.warning(self, "Ошибка", f"Ошибка загрузки программ: {exc}")
        self.combo_program.blockSignals(False)

    def _load_templates(self):
        """Load available templates into combobox."""
        self.combo_template.blockSignals(True)
        self.combo_template.clear()
        self.combo_template.addItem("-- Авто (по типу приказа) --", None)
        templates = self.generator.get_available_templates()
        for tpl in templates:
            self.combo_template.addItem(tpl, tpl)
        self.combo_template.blockSignals(False)

    def _fill_from_entry(self):
        """Populate form fields from the journal entry dict."""
        if not self.entry:
            return

        # Order type
        jtype = self.entry.get('journal_type', '')
        idx = self.combo_order_type.findData(jtype)
        if idx >= 0:
            self.combo_order_type.setCurrentIndex(idx)

        # Order number
        self.edit_order_number.setText(str(self.entry.get('order_number', '')))

        # Order date
        od = self.entry.get('order_date')
        if od:
            self.edit_order_date.setDate(QDate(od.year, od.month, od.day))
        else:
            self.edit_order_date.setDate(QDate.currentDate())

        # Title / Notes
        self.edit_title.setText(self.entry.get('title', ''))
        self.edit_notes.setText(self.entry.get('notes', '') or '')

        # Executor — try to split "position name" stored in journal
        executor_raw = self.entry.get('executor', '') or ''
        self._fill_executor_from_stored(executor_raw)

        # Program
        pid = self.entry.get('program_id')
        if pid:
            pidx = self.combo_program.findData(pid)
            if pidx >= 0:
                self.combo_program.setCurrentIndex(pidx)

        # Auto-select template by order type
        self._select_template_by_type(jtype)

        # Load listeners
        self._load_program_listeners()

    def _fill_executor_from_stored(self, executor_raw: str):
        """Try to split stored executor string into position + name."""
        if not executor_raw:
            return

        # Try to match against known positions in combo
        matched = False
        for i in range(self.combo_executor_position.count()):
            pos_text = self.combo_executor_position.itemText(i)
            if executor_raw.startswith(pos_text):
                self.combo_executor_position.setCurrentIndex(i)
                name_part = executor_raw[len(pos_text):].strip()
                self.edit_executor.setText(name_part)
                matched = True
                break

        if not matched:
            # Check if it contains typical keywords (факультет, кафедр, специалист, etc.)
            # to separate position from name
            keywords = [
                '(ФПК ', 'факультета ', 'кафедры ', 'отдела ', 'департамента ',
            ]
            for kw in keywords:
                idx = executor_raw.rfind(kw)
                if idx >= 0:
                    # Find the end of the position part (after closing parenthesis or keyword group)
                    paren_end = executor_raw.find(')', idx)
                    if paren_end >= 0:
                        pos_part = executor_raw[:paren_end + 1].strip()
                        name_part = executor_raw[paren_end + 1:].strip()
                    else:
                        pos_part = executor_raw[:idx + len(kw)].strip()
                        name_part = executor_raw[idx + len(kw):].strip()
                    self.combo_executor_position.setEditText(pos_part)
                    self.edit_executor.setText(name_part)
                    matched = True
                    break

        if not matched:
            # Fallback: put everything into name field
            self.edit_executor.setText(executor_raw)

    def _select_template_by_type(self, order_type: str):
        """Try to select the matching template for the order type."""
        tpl_name = self.generator.TEMPLATE_TYPES.get(order_type)
        if tpl_name:
            idx = self.combo_template.findData(tpl_name)
            if idx >= 0:
                self.combo_template.setCurrentIndex(idx)

    # ------------------------------------------------------------------
    # Template preview & editor
    # ------------------------------------------------------------------

    def _on_template_changed(self):
        """Called when template combo changes - update preview and editor."""
        self._update_template_preview()
        self._load_template_into_editor()

    def _update_template_preview(self):
        """Update the preview tab with the template preview."""
        tpl_name = self.combo_template.currentData()
        if not tpl_name:
            # Auto-select by order type
            order_type = self.combo_order_type.currentData()
            tpl_name = self.generator.TEMPLATE_TYPES.get(order_type)

        if not tpl_name:
            self.template_browser.setHtml(
                "<p style='color:gray;'>Выберите шаблон или тип приказа для предпросмотра</p>"
            )
            self.label_variables.setText("")
            return

        tpl_path = self.generator.templates_dir / tpl_name
        if not tpl_path.exists():
            self.template_browser.setHtml(
                f"<p style='color:orange;'>Шаблон не найден: {html_module.escape(tpl_name)}</p>"
                "<p>Поместите файл шаблона в папку <b>templates/</b></p>"
            )
            self.label_variables.setText("")
            return

        html_content = extract_template_html(str(tpl_path))
        self.template_browser.setHtml(html_content)

        # Show variable list
        variables = extract_template_variables(str(tpl_path))
        if variables:
            var_text = "Переменные: " + ", ".join(variables)
            self.label_variables.setText(var_text)
        else:
            self.label_variables.setText("Переменные не обнаружены в шаблоне")

    def _load_template_into_editor(self):
        """Load the text content of the template into the editor tab."""
        tpl_name = self._resolve_current_template_name()
        if not tpl_name:
            self.template_editor.clear()
            return

        tpl_path = self.generator.templates_dir / tpl_name
        if not tpl_path.exists():
            self.template_editor.setPlainText(f"(Шаблон не найден: {tpl_name})")
            return

        text = extract_template_plain_text(str(tpl_path))
        self.template_editor.setPlainText(text)

    def _resolve_current_template_name(self) -> Optional[str]:
        """Get current template name (from combo or auto by order type)."""
        tpl_name = self.combo_template.currentData()
        if not tpl_name:
            order_type = self.combo_order_type.currentData()
            tpl_name = self.generator.TEMPLATE_TYPES.get(order_type)
        return tpl_name

    # ------------------------------------------------------------------
    # Variable categories & quick insert
    # ------------------------------------------------------------------

    def _populate_var_categories(self):
        """Populate the variable category combobox."""
        self.combo_var_category.blockSignals(True)
        self.combo_var_category.clear()
        for cat_name, _ in VARIABLE_REFERENCE.items():
            self.combo_var_category.addItem(cat_name, cat_name)
        self.combo_var_category.blockSignals(False)
        self._on_var_category_changed()

    def _on_var_category_changed(self):
        """Populate variable names for selected category."""
        self.combo_var_name.blockSignals(True)
        self.combo_var_name.clear()
        cat = self.combo_var_category.currentData()
        if cat and cat in VARIABLE_REFERENCE:
            for var_info in VARIABLE_REFERENCE[cat]:
                display = f"{var_info['var']}  —  {var_info['desc']}"
                self.combo_var_name.addItem(display, var_info)
        self.combo_var_name.blockSignals(False)
        self._on_var_selected()

    def _on_var_selected(self):
        """Show description of the selected variable."""
        var_info = self.combo_var_name.currentData()
        if var_info:
            self.label_var_desc.setText(
                f"<b>{var_info['var']}</b> → {var_info['desc']}"
                f"<br>Пример: <i>{var_info.get('example', '')}</i>"
            )
        else:
            self.label_var_desc.setText("")

    def _copy_variable(self):
        """Copy selected variable to clipboard."""
        var_info = self.combo_var_name.currentData()
        if not var_info:
            return
        text = "{{ " + var_info['var'] + " }}"
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(
            self, "Скопировано",
            f"Переменная скопирована в буфер обмена:\n{text}"
        )

    def _insert_variable(self):
        """Insert selected variable at cursor position in the editor."""
        var_info = self.combo_var_name.currentData()
        if not var_info:
            return
        text = "{{ " + var_info['var'] + " }}"
        # Switch to editor tab
        self.left_tabs.setCurrentIndex(1)
        cursor = self.template_editor.textCursor()
        cursor.insertText(text)
        self.template_editor.setFocus()

    def _add_variable_at_cursor(self):
        """Prompt for a custom variable name and insert at cursor."""
        name, ok = QInputDialog.getText(
            self, "Добавить переменную",
            "Имя переменной (например: my_variable):",
        )
        if ok and name.strip():
            var_name = name.strip().replace(' ', '_')
            text = "{{ " + var_name + " }}"
            cursor = self.template_editor.textCursor()
            cursor.insertText(text)
            self.template_editor.setFocus()

    def _replace_variable(self):
        """Find and replace a variable in the template editor text."""
        dialog = _VariableReplaceDialog(self, self.template_editor.toPlainText())
        if dialog.exec_():
            old_var, new_var = dialog.get_result()
            if old_var and new_var:
                text = self.template_editor.toPlainText()
                # Replace {{ old }} with {{ new }}
                old_pattern = "{{ " + old_var + " }}"
                new_pattern = "{{ " + new_var + " }}"
                if old_pattern not in text:
                    # Try without spaces
                    old_pattern = "{{" + old_var + "}}"
                count = text.count(old_pattern)
                if count == 0:
                    QMessageBox.warning(
                        self, "Не найдено",
                        f"Переменная «{old_pattern}» не найдена в тексте шаблона."
                    )
                    return
                new_text = text.replace(old_pattern, new_pattern)
                self.template_editor.setPlainText(new_text)
                QMessageBox.information(
                    self, "Замена выполнена",
                    f"Заменено {count} вхождений:\n{old_pattern} → {new_pattern}"
                )

    def _save_template(self):
        """Save the edited text back into the .docx template."""
        tpl_name = self._resolve_current_template_name()
        if not tpl_name:
            QMessageBox.warning(self, "Ошибка", "Не выбран шаблон для сохранения")
            return

        tpl_path = self.generator.templates_dir / tpl_name
        if not tpl_path.exists():
            QMessageBox.warning(self, "Ошибка", f"Файл шаблона не найден: {tpl_name}")
            return

        new_text = self.template_editor.toPlainText()
        if not new_text.strip():
            QMessageBox.warning(self, "Ошибка", "Текст шаблона пуст")
            return

        reply = QMessageBox.question(
            self, "Сохранение шаблона",
            f"Сохранить изменения в шаблон «{tpl_name}»?\n\n"
            "⚠️ Будет создана резервная копия перед сохранением.\n"
            "Внимание: форматирование (шрифты, стили) сохраняется,\n"
            "изменяется только текстовое содержимое.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes,
        )
        if reply != QMessageBox.Yes:
            return

        try:
            save_text_to_template(str(tpl_path), new_text)
            QMessageBox.information(self, "Успех", f"Шаблон «{tpl_name}» сохранён!")
            # Refresh preview
            self._update_template_preview()
        except Exception as exc:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить шаблон:\n{exc}")

    @staticmethod
    def _build_variable_reference_html() -> str:
        """Build HTML for the variable reference tab."""
        parts = [
            "<style>",
            "body { font-family: -apple-system, 'Segoe UI', sans-serif; font-size: 13px; line-height: 1.5; }",
            "h2 { color: #1976D2; margin-top: 5px; }",
            "h3 { color: #1565C0; margin-top: 18px; background: #E3F2FD; padding: 6px 10px; border-radius: 5px; }",
            "table { border-collapse: collapse; width: 100%; margin: 6px 0; }",
            "th { background: #1976D2; color: white; padding: 6px 8px; text-align: left; font-size: 12px; }",
            "td { padding: 5px 8px; border-bottom: 1px solid #E0E0E0; font-size: 12px; }",
            "tr:hover { background: #F5F5F5; }",
            "code { background: #ECEFF1; padding: 1px 5px; border-radius: 3px; font-family: 'Courier New', monospace; color: #C62828; font-size: 12px; }",
            ".note { background: #FFF8E1; border-left: 3px solid #FFC107; padding: 8px; margin: 8px 0; border-radius: 3px; font-size: 12px; }",
            ".tip { background: #E8F5E9; border-left: 3px solid #4CAF50; padding: 8px; margin: 8px 0; border-radius: 3px; font-size: 12px; }",
            "</style>",
            "<h2>📚 Полный справочник переменных</h2>",
        ]

        for cat_name, variables in VARIABLE_REFERENCE.items():
            parts.append(f"<h3>{cat_name}</h3>")
            parts.append("<table><tr><th>Переменная</th><th>Описание</th><th>Пример</th></tr>")
            for v in variables:
                var_code = html_module.escape(v['var'])
                desc = html_module.escape(v['desc'])
                example = html_module.escape(v.get('example', ''))
                parts.append(
                    f"<tr><td><code>{{{{ {var_code} }}}}</code></td>"
                    f"<td>{desc}</td><td><i>{example}</i></td></tr>"
                )
            parts.append("</table>")

        # Add usage tips
        parts.append("""
        <h3>💡 Советы по использованию</h3>
        <div class="tip">
        <b>Для списка слушателей (таблица)</b> используйте переменные с префиксом
        <code>listener.</code> — они работают внутри цикла и подставляются для каждого слушателя.
        </div>
        <div class="tip">
        <b>Для индивидуальных документов</b> (справки, удостоверения) используйте переменные
        без префикса — <code>full_name</code>, <code>position</code> и т.д.
        </div>
        <div class="note">
        <b>Склонение ФИО:</b> суффиксы <code>_genitive</code> (род.), <code>_dative</code> (дат.),
        <code>_accusative</code> (вин.), <code>_instrumental</code> (тв.), <code>_prepositional</code> (пр.)
        работают для <code>full_name</code>, <code>last_name</code>, <code>first_name</code>, <code>middle_name</code>.
        </div>
        <div class="note">
        <b>Переменные цикла:</b> <code>{{ loop.index }}</code> — номер по порядку (1, 2, 3...).
        Доступен только внутри таблицы с циклом <code>{%tr for listener in listeners %}</code>.
        </div>
        """)

        return "\n".join(parts)

    def _open_templates_folder(self):
        """Open the templates folder in file manager."""
        folder = self.generator.templates_dir
        folder.mkdir(parents=True, exist_ok=True)
        try:
            if sys.platform == 'darwin':
                subprocess.run(['open', str(folder)])
            elif sys.platform == 'win32':
                subprocess.run(['explorer', str(folder)])
            else:
                subprocess.run(['xdg-open', str(folder)])
        except Exception as exc:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть папку: {exc}")

    def _fix_all_templates(self):
        """Fix all templates by adding {% for listener in listeners %} loop."""
        reply = QMessageBox.question(
            self, "Исправление шаблонов",
            "Будет выполнено автоматическое исправление всех шаблонов:\n\n"
            "• Добавление цикла {% for listener in listeners %} в строки таблиц,\n"
            "  содержащие переменные listener/loop.\n\n"
            "Продолжить?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply != QMessageBox.Yes:
            return

        QApplication.processEvents()
        results = self.generator.fix_all_templates()

        if not results:
            QMessageBox.information(
                self, "Информация",
                "В папке templates/ нет шаблонов .docx"
            )
            return

        fixed = [r for r in results if r['status'] == 'fixed']
        already = [r for r in results if r['status'] == 'already_ok']
        no_vars = [r for r in results if r['status'] == 'no_vars']
        errors = [r for r in results if r['status'] == 'error']

        lines = []
        if fixed:
            lines.append(f"✅ Исправлено: {len(fixed)}")
            for r in fixed:
                lines.append(f"   • {r['name']}")
        if already:
            lines.append(f"\nℹ️ Уже содержат цикл: {len(already)}")
            for r in already:
                lines.append(f"   • {r['name']}")
        if no_vars:
            lines.append(f"\n⚠ Без переменных listener/loop: {len(no_vars)}")
            for r in no_vars:
                lines.append(f"   • {r['name']}")
        if errors:
            lines.append(f"\n❌ Ошибки: {len(errors)}")
            for r in errors:
                lines.append(f"   • {r['name']}: {r['message']}")

        msg = "\n".join(lines)
        QMessageBox.information(self, "Результат исправления шаблонов", msg)

        # Refresh templates list and preview
        self._load_templates()
        self._update_template_preview()

    def _on_type_changed(self):
        """When order type changes, update template and number hint."""
        order_type = self.combo_order_type.currentData()
        if order_type:
            self._select_template_by_type(order_type)
            self._update_template_preview()

    # ------------------------------------------------------------------
    # Program & listeners
    # ------------------------------------------------------------------

    def _on_program_changed(self):
        """When program changes, reload listeners and auto-fill fields."""
        self._load_program_listeners()
        prog_id = self.combo_program.currentData()
        if prog_id:
            self._fill_fields_from_program(prog_id)

    def _load_program_listeners(self):
        """Load listeners for the selected program."""
        self.list_listeners.clear()
        self._program_listeners.clear()

        prog_id = self.combo_program.currentData()
        if not prog_id:
            self._update_listener_count()
            return

        try:
            with DatabaseSession() as session:
                assocs = session.query(ProgramListener).filter(
                    ProgramListener.program_id == prog_id
                ).order_by(ProgramListener.order_number).all()

                for a in assocs:
                    listener = session.query(Listener).get(a.listener_id)
                    if listener:
                        ld = {
                            'id': listener.id,
                            'full_name': listener.full_name,
                            'last_name': listener.last_name,
                            'first_name': listener.first_name,
                            'middle_name': listener.middle_name,
                            'position': listener.position,
                            'workplace': listener.workplace,
                            'court_name': listener.workplace,
                            'region': listener.region,
                            'notes': listener.notes,
                        }
                        self._program_listeners.append(ld)
                        self._add_listener_item(ld, checked=True)
        except Exception as exc:
            QMessageBox.warning(self, "Ошибка", f"Ошибка загрузки слушателей: {exc}")

        # Re-add extra listeners (manually added)
        for ld in self._extra_listeners:
            if not any(pl['id'] == ld['id'] for pl in self._program_listeners):
                self._add_listener_item(ld, checked=True, extra=True)

        self._update_listener_count()

    def _add_listener_item(self, ld: Dict[str, Any], checked: bool = True, extra: bool = False):
        """Add a listener item to the list widget."""
        name = ld.get('full_name', '')
        position = ld.get('position') or 'Должность не указана'
        workplace = ld.get('workplace') or 'Место работы не указано'
        text = f"{name} — {position} ({workplace})"
        if extra:
            text = f"[+] {text}"

        item = QListWidgetItem(text)
        item.setData(Qt.UserRole, ld['id'])
        item.setData(Qt.UserRole + 1, ld)  # store full dict
        item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
        item.setCheckState(Qt.Checked if checked else Qt.Unchecked)

        if extra:
            item.setForeground(QColor("#1565C0"))

        self.list_listeners.addItem(item)

    def _fill_fields_from_program(self, prog_id: int):
        """Auto-fill education form/format and dates from program data."""
        program_data = None
        for p in self._all_programs:
            if p['id'] == prog_id:
                program_data = p
                break
        if not program_data:
            return

        # Education form
        edu_form = (program_data.get('education_form') or '').strip()
        if edu_form:
            idx = self.combo_education_form.findData(edu_form)
            if idx >= 0:
                self.combo_education_form.setCurrentIndex(idx)

        # Education format
        edu_format = (program_data.get('education_format') or '').strip()
        if edu_format:
            idx = self.combo_education_format.findData(edu_format)
            if idx >= 0:
                self.combo_education_format.setCurrentIndex(idx)
            else:
                self.combo_education_format.setEditText(edu_format)

        # Hours from volume
        volume = (program_data.get('program_volume') or '').strip()
        if volume:
            match = re.search(r'(\d+)', volume)
            if match:
                self.spin_hours.setValue(int(match.group(1)))

        # Training period -> dates
        period = (program_data.get('training_period') or '').strip()
        if period:
            self._parse_training_period(period)

    def _parse_training_period(self, period: str):
        """Parse training period string into start/end dates."""
        pattern = r'(\d{2}[./]\d{2}[./]\d{4})\s*[-–—]\s*(\d{2}[./]\d{2}[./]\d{4})'
        match = re.search(pattern, period)
        if not match:
            pattern = r'с?\s*(\d{2}[./]\d{2}[./]\d{4})\s*(?:по|до)\s*(\d{2}[./]\d{2}[./]\d{4})'
            match = re.search(pattern, period)
        if match:
            try:
                d1 = datetime.strptime(match.group(1).replace('/', '.'), '%d.%m.%Y').date()
                d2 = datetime.strptime(match.group(2).replace('/', '.'), '%d.%m.%Y').date()
                self.edit_start_date.setDate(QDate(d1.year, d1.month, d1.day))
                self.edit_end_date.setDate(QDate(d2.year, d2.month, d2.day))
            except ValueError:
                pass

    # ------------------------------------------------------------------
    # Listener actions
    # ------------------------------------------------------------------

    def _select_all_listeners(self):
        for i in range(self.list_listeners.count()):
            self.list_listeners.item(i).setCheckState(Qt.Checked)
        self._update_listener_count()

    def _deselect_all_listeners(self):
        for i in range(self.list_listeners.count()):
            self.list_listeners.item(i).setCheckState(Qt.Unchecked)
        self._update_listener_count()

    def _add_listener(self):
        """Show dialog to pick a listener not yet in the list."""
        existing_ids = set()
        for i in range(self.list_listeners.count()):
            lid = self.list_listeners.item(i).data(Qt.UserRole)
            existing_ids.add(lid)

        dialog = _ListenerPickerDialog(self, exclude_ids=existing_ids)
        if dialog.exec_() and dialog.selected_listeners:
            for ld in dialog.selected_listeners:
                self._extra_listeners.append(ld)
                self._add_listener_item(ld, checked=True, extra=True)
            self._update_listener_count()

    def _remove_listener(self):
        """Remove (uncheck) selected listeners in the list."""
        selected_items = self.list_listeners.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "Информация", "Выберите слушателей для удаления из списка")
            return

        for item in selected_items:
            item.setCheckState(Qt.Unchecked)
        self._update_listener_count()

    def _update_listener_count(self):
        count = sum(
            1 for i in range(self.list_listeners.count())
            if self.list_listeners.item(i).checkState() == Qt.Checked
        )
        self.label_listeners_count.setText(f"Выбрано: {count}")

    def _get_selected_listeners(self) -> List[Dict[str, Any]]:
        """Get list of checked listener dicts."""
        selected = []
        for i in range(self.list_listeners.count()):
            item = self.list_listeners.item(i)
            if item.checkState() == Qt.Checked:
                ld = item.data(Qt.UserRole + 1)
                if ld:
                    selected.append(ld)
        return selected

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------

    def _open_document(self):
        """Open the current document file."""
        doc_path = self.entry.get('document_path')
        if not doc_path:
            QMessageBox.information(self, "Информация", "Документ не привязан к этой записи")
            return
        p = Path(doc_path)
        if not p.exists():
            QMessageBox.warning(self, "Предупреждение", f"Файл не найден:\n{doc_path}")
            return
        try:
            if sys.platform == 'darwin':
                subprocess.run(['open', str(p)])
            elif sys.platform == 'win32':
                subprocess.run(['start', '', str(p)], shell=True)
            else:
                subprocess.run(['xdg-open', str(p)])
        except Exception as exc:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть файл: {exc}")

    def _open_output_folder(self):
        """Open the output folder in file manager."""
        output_path = self.generator.output_dir
        try:
            if sys.platform == 'darwin':
                subprocess.run(['open', str(output_path)])
            elif sys.platform == 'win32':
                subprocess.run(['explorer', str(output_path)])
            else:
                subprocess.run(['xdg-open', str(output_path)])
        except Exception:
            pass

    def _save_journal_entry(self):
        """Save changes to the journal entry (metadata only, no regeneration)."""
        entry_id = self.entry.get('id')
        if not entry_id:
            QMessageBox.warning(self, "Ошибка", "Запись журнала не найдена")
            return

        num_text = self.edit_order_number.text().strip()
        if not num_text.isdigit():
            QMessageBox.warning(self, "Ошибка", "Номер приказа должен быть числом")
            return

        executor_position = self.combo_executor_position.currentText().strip()
        executor_name = self.edit_executor.text().strip()
        executor_full = f"{executor_position} {executor_name}".strip() if executor_position else executor_name

        data = {
            'journal_type': self.combo_order_type.currentData(),
            'order_number': int(num_text),
            'order_date': self.edit_order_date.date().toPyDate(),
            'title': self.edit_title.text().strip(),
            'executor': executor_full,
            'program_id': self.combo_program.currentData(),
            'notes': self.edit_notes.text().strip() or None,
        }

        # Resolve program name
        pid = data['program_id']
        if pid:
            for p in self._all_programs:
                if p['id'] == pid:
                    data['program_name'] = p.get('program_name') or p.get('program_short_name', '')
                    break

        try:
            self.journal_service.update_order_entry(entry_id, **data)
            QMessageBox.information(self, "Успех", "Запись журнала обновлена")
        except Exception as exc:
            QMessageBox.critical(self, "Ошибка", f"Не удалось обновить запись: {exc}")

    def _regenerate_document(self):
        """Regenerate the document with current parameters."""
        # Validate
        order_number = self.edit_order_number.text().strip()
        if not order_number:
            QMessageBox.warning(self, "Предупреждение", "Введите номер приказа")
            return
        if not order_number.isdigit():
            QMessageBox.warning(self, "Предупреждение", "Номер приказа должен быть числом")
            return

        selected_listeners = self._get_selected_listeners()
        if not selected_listeners:
            QMessageBox.warning(self, "Предупреждение", "Выберите хотя бы одного слушателя")
            return

        reply = QMessageBox.question(
            self, "Подтверждение",
            f"Перегенерировать документ с {len(selected_listeners)} слушателями?\n\n"
            "Старый документ будет заменён новым.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes,
        )
        if reply != QMessageBox.Yes:
            return

        order_type = self.combo_order_type.currentData()
        order_type_label = ORDER_TYPE_LABELS.get(order_type, '')
        order_date = self.edit_order_date.date().toPyDate()
        start_date = self.edit_start_date.date().toPyDate()
        end_date = self.edit_end_date.date().toPyDate()

        order_datetime = datetime.combine(order_date, datetime.min.time())
        start_datetime = datetime.combine(start_date, datetime.min.time())
        end_datetime = datetime.combine(end_date, datetime.min.time())

        prog_id = self.combo_program.currentData()
        program_name = ""
        training_period = ""
        training_duration = ""
        for p in self._all_programs:
            if p['id'] == prog_id:
                program_name = p.get('program_name') or p.get('program_short_name', '')
                training_period = p.get('training_period') or ''
                training_duration = p.get('training_duration') or ''
                break

        education_form = self.combo_education_form.currentData() or ''
        education_format = self.combo_education_format.currentText().strip()
        program_type = self.combo_program_type.currentData() or ''
        executor_position = self.combo_executor_position.currentText().strip()
        executor_name = self.edit_executor.text().strip()
        executor_full = f"{executor_position} {executor_name}".strip() if executor_position else executor_name

        # Determine template
        selected_template = self.combo_template.currentData()  # None = auto

        QApplication.processEvents()
        self.btn_regenerate.setEnabled(False)
        self.btn_regenerate.setText("⏳ Генерация...")

        try:
            file_path = self.generator.generate_order(
                order_type=order_type,
                listeners_data=selected_listeners,
                order_number=order_number,
                order_date=order_datetime,
                program_name=program_name,
                start_date=start_datetime,
                end_date=end_datetime,
                hours=self.spin_hours.value(),
                education_form=education_form,
                education_format=education_format,
                template_name=selected_template,
                custom_context={
                    'order_type_label': order_type_label,
                    'program_type': program_type,
                    'program_id': prog_id,
                    'executor_position': executor_position,
                    'executor_name': executor_name,
                    'executor_full': executor_full,
                    'training_period': training_period,
                    'training_duration': training_duration,
                },
            )

            # Update journal entry with new document path
            entry_id = self.entry.get('id')
            if entry_id:
                update_data = {
                    'journal_type': order_type,
                    'order_number': int(order_number),
                    'order_date': order_date,
                    'title': self.edit_title.text().strip() or f"Приказ «{order_type_label}» №{order_number}",
                    'executor': executor_full,
                    'program_id': prog_id,
                    'program_name': program_name,
                    'notes': f"Слушателей: {len(selected_listeners)}",
                    'document_path': file_path,
                }
                try:
                    self.journal_service.update_order_entry(entry_id, **update_data)
                except Exception:
                    pass  # non-critical

            # Update our entry reference
            self.entry['document_path'] = file_path

            QMessageBox.information(
                self, "Успех",
                f"Документ перегенерирован!\n\n"
                f"Файл: {Path(file_path).name}\n"
                f"Слушателей: {len(selected_listeners)}"
            )

            if self.check_open_folder.isChecked():
                self._open_output_folder()

        except FileNotFoundError as exc:
            QMessageBox.critical(
                self, "Ошибка",
                f"Шаблон не найден!\n\n{exc}\n\n"
                "Убедитесь, что файлы шаблонов находятся в папке templates/"
            )
        except Exception as exc:
            QMessageBox.critical(self, "Ошибка", f"Ошибка генерации: {exc}")
        finally:
            self.btn_regenerate.setEnabled(True)
            self.btn_regenerate.setText("🔄 Перегенерировать документ")


# ---------------------------------------------------------------------------
# Helper dialog: find & replace variable
# ---------------------------------------------------------------------------

class _VariableReplaceDialog(QDialog):
    """Dialog for finding and replacing a variable name in template text."""

    def __init__(self, parent=None, template_text: str = ""):
        super().__init__(parent)
        self.template_text = template_text
        self._result = (None, None)
        self._setup_ui()

    def _setup_ui(self):
        self.setWindowTitle("Заменить переменную")
        self.setMinimumWidth(500)
        self.setModal(True)

        layout = QVBoxLayout(self)

        # Existing variables
        existing_vars = set()
        for m in re.finditer(r'\{\{\s*([\w.]+)\s*\}\}', self.template_text):
            existing_vars.add(m.group(1))
        existing_vars_sorted = sorted(existing_vars)

        form = QFormLayout()

        self.combo_old = QComboBox()
        self.combo_old.setEditable(True)
        for v in existing_vars_sorted:
            self.combo_old.addItem(v)
        form.addRow("Найти переменную:", self.combo_old)

        # New variable: either pick from reference or type custom
        self.combo_new = QComboBox()
        self.combo_new.setEditable(True)
        self.combo_new.setMinimumWidth(300)
        # Add all known variables
        for cat_name, variables in VARIABLE_REFERENCE.items():
            for v in variables:
                self.combo_new.addItem(f"{v['var']}  ({v['desc']})", v['var'])
        form.addRow("Заменить на:", self.combo_new)

        layout.addLayout(form)

        if existing_vars_sorted:
            info = QLabel(f"Найдено {len(existing_vars_sorted)} переменных в шаблоне")
            info.setStyleSheet("color: gray; font-size: 11px; padding: 4px;")
            layout.addWidget(info)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_replace = QPushButton("Заменить")
        btn_replace.setDefault(True)
        btn_replace.clicked.connect(self._on_replace)
        btn_cancel = QPushButton("Отмена")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_replace)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def _on_replace(self):
        old_text = self.combo_old.currentText().strip()
        new_data = self.combo_new.currentData()
        new_text = new_data if new_data else self.combo_new.currentText().strip()
        # Clean up: remove description in parentheses if user didn't change
        if '(' in new_text:
            new_text = new_text.split('(')[0].strip()
        if not old_text or not new_text:
            QMessageBox.warning(self, "Ошибка", "Укажите обе переменные")
            return
        self._result = (old_text, new_text)
        self.accept()

    def get_result(self):
        return self._result


# ---------------------------------------------------------------------------
# Helper dialog: pick listeners from database
# ---------------------------------------------------------------------------

class _ListenerPickerDialog(QDialog):
    """Small dialog to pick listeners from the full database list."""

    def __init__(self, parent=None, exclude_ids: set = None):
        super().__init__(parent)
        self.exclude_ids = exclude_ids or set()
        self.selected_listeners: List[Dict[str, Any]] = []
        self._all_listeners: List[Dict[str, Any]] = []
        self._setup_ui()
        self._load_listeners()

    def _setup_ui(self):
        self.setWindowTitle("Добавить слушателя")
        self.setMinimumSize(600, 500)
        self.setModal(True)

        layout = QVBoxLayout(self)

        # Search
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Поиск:"))
        self.edit_search = QLineEdit()
        self.edit_search.setPlaceholderText("Введите ФИО для поиска...")
        self.edit_search.setClearButtonEnabled(True)
        self.edit_search.textChanged.connect(self._filter_list)
        search_layout.addWidget(self.edit_search)
        layout.addLayout(search_layout)

        # List
        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        layout.addWidget(self.list_widget)

        self.label_info = QLabel("")
        self.label_info.setStyleSheet("color: gray;")
        layout.addWidget(self.label_info)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_add = QPushButton("Добавить выбранных")
        btn_add.setDefault(True)
        btn_add.clicked.connect(self._on_add)
        btn_cancel = QPushButton("Отмена")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_add)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def _load_listeners(self):
        """Load all listeners from DB, excluding already-added ones."""
        try:
            with DatabaseSession() as session:
                listeners = session.query(Listener).order_by(Listener.last_name).all()
                for l in listeners:
                    if l.id in self.exclude_ids:
                        continue
                    ld = {
                        'id': l.id,
                        'full_name': l.full_name,
                        'last_name': l.last_name,
                        'first_name': l.first_name,
                        'middle_name': l.middle_name,
                        'position': l.position,
                        'workplace': l.workplace,
                        'court_name': l.workplace,
                        'region': l.region,
                        'notes': l.notes,
                    }
                    self._all_listeners.append(ld)
        except Exception:
            pass

        self._populate_list(self._all_listeners)

    def _populate_list(self, listeners: List[Dict[str, Any]]):
        self.list_widget.clear()
        for ld in listeners:
            text = (
                f"{ld.get('full_name', '')} — "
                f"{ld.get('position') or 'Должность не указана'} "
                f"({ld.get('workplace') or 'Место работы не указано'})"
            )
            item = QListWidgetItem(text)
            item.setData(Qt.UserRole, ld)
            self.list_widget.addItem(item)
        self.label_info.setText(f"Доступно: {len(listeners)}")

    def _filter_list(self, text: str):
        search = text.strip().lower()
        if not search:
            self._populate_list(self._all_listeners)
            return
        filtered = [
            ld for ld in self._all_listeners
            if search in (ld.get('full_name', '') or '').lower()
            or search in (ld.get('position', '') or '').lower()
            or search in (ld.get('workplace', '') or '').lower()
        ]
        self._populate_list(filtered)

    def _on_add(self):
        selected = self.list_widget.selectedItems()
        if not selected:
            QMessageBox.information(self, "Информация", "Выберите слушателей из списка")
            return
        self.selected_listeners = [item.data(Qt.UserRole) for item in selected]
        self.accept()
