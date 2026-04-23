"""
Document generator service for creating Word documents from templates.
Uses docxtpl (Jinja2-based) for template processing.
"""

import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docxtpl import DocxTemplate

from ..database.models import Listener, Program, ProgramListener, ContractJournal
from ..database.connection import DatabaseSession
from .declension import DeclensionService, get_declension_service
from .document_registration import DocumentRegistrationService


def get_app_dir() -> Path:
    """Get application directory (works for both dev and compiled)."""
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent
    else:
        return Path(__file__).parent.parent.parent


class DocumentGenerator:
    """
    Generate Word documents from templates in templates/ folder.
    
    Uses Jinja2 templating syntax via docxtpl library.
    Supports name declension for Russian cases.
    """
    
    # Predefined template types
    TEMPLATE_TYPES = {
        'enrollment': 'О_зачислении.docx',
        'admission': 'О_допуске.docx',
        'graduation': 'Об_отчислении.docx',
        'thesis_topics': 'Об_утверждении_тем.docx',
        'internship': 'О_направлении_на_стажировку.docx',
        'theory_exam': 'О_допуске_к_экзамену.docx',
        'theory_exam_retake': 'О_допуске_к_повторному_экзамену.docx',
    }
    
    def __init__(
        self, 
        templates_dir: Union[str, Path] = None,
        output_dir: Union[str, Path] = None
    ):
        """
        Initialize the document generator.
        
        Args:
            templates_dir: Directory containing Word templates
            output_dir: Directory for generated documents
        """
        app_dir = get_app_dir()
        self.templates_dir = Path(templates_dir) if templates_dir else app_dir / "templates"
        self.output_dir = Path(output_dir) if output_dir else app_dir / "output"
        
        # Ensure directories exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize declension service
        self.declension = get_declension_service()
        
        # Initialize document registration service
        self.registration_service = DocumentRegistrationService()
        self.enable_registration = True  # Can be disabled if needed
    
    def get_available_templates(self) -> List[str]:
        """
        Get list of available template files.
        
        Returns:
            List of template filenames
        """
        if not self.templates_dir.exists():
            return []
        
        seen: set = set()
        templates: List[str] = []
        for pattern in ('*.docx', '*.DOCX'):
            for t in self.templates_dir.glob(pattern):
                name = t.name
                key = name.lower()
                if key in seen:
                    continue
                if name.startswith('~$') or name.startswith('._'):
                    continue
                seen.add(key)
                templates.append(name)

        return sorted(templates)
    
    def generate_for_listener(
        self,
        listener: Listener,
        program: Optional[Program],
        template_name: str,
        order_number: int = 1,
        custom_context: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate a document for a single listener.
        
        Args:
            listener: Listener instance
            program: Program instance (optional)
            template_name: Name of the template file
            order_number: Sequence number for the listener
            custom_context: Additional template variables
            output_filename: Custom output filename
        
        Returns:
            Path to the generated document
        """
        # Resolve template path
        template_path = self._resolve_template_path(template_name)
        
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        
        # Load template
        doc = DocxTemplate(template_path)
        
        # Prepare context
        context = self._prepare_context(listener, program, order_number)
        
        # Add custom context
        if custom_context:
            context.update(custom_context)
        
        # Render document
        doc.render(context)
        
        # Generate output filename
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = self._sanitize_filename(listener.last_name)
            output_name = f"{template_path.stem}_{safe_name}_{timestamp}.docx"
            output_path = self.output_dir / output_name
        
        # Save document
        doc.save(output_path)
        
        # Register document with automatic numbering
        if self.enable_registration:
            doc_type = self._determine_document_type(template_name)
            listener_name = f"{listener.last_name} {listener.first_name}"
            if listener.middle_name:
                listener_name += f" {listener.middle_name}"
            
            self.registration_service.register_document(
                document_type=doc_type,
                document_path=str(output_path),
                listener_name=listener_name,
                listener_id=listener.id,
                program_id=program.id if program else None
            )
        
        return str(output_path)
    
    def generate_for_listener_dict(
        self,
        listener_data: Dict[str, Any],
        program_data: Optional[Dict[str, Any]],
        template_name: str,
        order_number: int = 1,
        custom_context: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate a document for a listener using dictionary data.
        
        Args:
            listener_data: Listener data as dictionary
            program_data: Program data as dictionary (optional)
            template_name: Name of the template file
            order_number: Sequence number for the listener
            custom_context: Additional template variables
            output_filename: Custom output filename
        
        Returns:
            Path to the generated document
        """
        # Resolve template path
        template_path = self._resolve_template_path(template_name)
        
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        
        # Load template
        doc = DocxTemplate(template_path)
        
        # Prepare context from dictionaries
        context = self._prepare_context_from_dict(listener_data, program_data, order_number)
        
        # Add custom context
        if custom_context:
            context.update(custom_context)
        
        # Render document
        doc.render(context)
        
        # Generate output filename
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = self._sanitize_filename(listener_data.get('last_name', 'listener'))
            output_name = f"{template_path.stem}_{safe_name}_{timestamp}.docx"
            output_path = self.output_dir / output_name
        
        # Save document
        doc.save(output_path)
        
        # Register document with automatic numbering
        if self.enable_registration:
            doc_type = self._determine_document_type(template_name)
            listener_name = f"{listener_data.get('last_name', '')} {listener_data.get('first_name', '')}"
            if listener_data.get('middle_name'):
                listener_name += f" {listener_data.get('middle_name')}"
            
            self.registration_service.register_document(
                document_type=doc_type,
                document_path=str(output_path),
                listener_name=listener_name.strip(),
                listener_id=listener_data.get('id'),
                program_id=program_data.get('id') if program_data else None
            )
        
        return str(output_path)
    
    def _prepare_context_from_dict(
        self,
        listener_data: Dict[str, Any],
        program_data: Optional[Dict[str, Any]],
        order_number: int
    ) -> Dict[str, Any]:
        """
        Prepare complete template context from dictionaries.
        """
        context = {}
        
        # Get all name declensions
        context.update(self.declension.get_all_declensions(
            listener_data.get('last_name', ''),
            listener_data.get('first_name', ''),
            listener_data.get('middle_name')
        ))
        
        # Add other listener fields
        context['order_number'] = order_number
        context['position'] = listener_data.get('position') or ''
        context['workplace'] = listener_data.get('workplace') or ''
        context['court_name'] = listener_data.get('workplace') or ''  # Alias for templates
        context['region'] = listener_data.get('region') or ''
        context['notes'] = listener_data.get('notes') or ''
        
        # Contact information
        context['mobile_phone'] = listener_data.get('mobile_phone') or ''
        context['work_phone'] = listener_data.get('work_phone') or ''
        context['email'] = listener_data.get('email') or ''
        
        # Birth date
        birth_date = listener_data.get('birth_date')
        if birth_date:
            if hasattr(birth_date, 'strftime'):
                context['birth_date'] = birth_date.strftime("%d.%m.%Y")
            else:
                context['birth_date'] = str(birth_date)
        else:
            context['birth_date'] = ''
        
        # Passport data
        context['passport_series_number'] = listener_data.get('passport_series_number') or ''
        passport_issue_date = listener_data.get('passport_issue_date')
        if passport_issue_date:
            if hasattr(passport_issue_date, 'strftime'):
                context['passport_issue_date'] = passport_issue_date.strftime("%d.%m.%Y")
            else:
                context['passport_issue_date'] = str(passport_issue_date)
        else:
            context['passport_issue_date'] = ''
        context['passport_issued_by'] = listener_data.get('passport_issued_by') or ''
        context['passport_department_code'] = listener_data.get('passport_department_code') or ''
        
        # Addresses
        context['registration_address'] = listener_data.get('registration_address') or ''
        context['actual_address'] = listener_data.get('actual_address') or ''
        
        # Identifiers
        context['snils'] = listener_data.get('snils') or ''
        context['inn'] = listener_data.get('inn') or ''
        
        # Personal data consent
        context['personal_data_consent'] = 'Да' if listener_data.get('personal_data_consent') else 'Нет'
        
        # Add program context
        if program_data:
            context['program_name'] = program_data.get('program_name') or ''
            context['program_short_name'] = program_data.get('program_short_name') or ''
            basis = program_data.get('training_basis') or ''
            basis_instr = self._decline_phrase(basis, 'instrumental')
            context['training_basis'] = basis
            context['training_basis_instrumental'] = basis_instr
            context['training_basis_phrase'] = f"в соответствии с {basis_instr}" if basis_instr else ''
            context['training_period'] = program_data.get('training_period') or ''
            context['program_volume'] = program_data.get('program_volume') or ''
            context['education_form'] = program_data.get('education_form') or ''
            context['education_format'] = program_data.get('education_format') or ''
            context['listener_category'] = program_data.get('listener_category') or ''
            context['expulsion_date'] = program_data.get('formatted_expulsion_date') or ''
        
        # Add service context
        context.update(self._prepare_service_context())
        
        return context

    def generate_order(
        self,
        order_type: str,
        listeners_data: List[Dict[str, Any]],
        order_number: str,
        order_date: datetime,
        program_name: str = "",
        stream_name: str = "",
        start_date: Optional[datetime] = None,
        end_date: Optional[datetime] = None,
        contract_date: str = "",
        contract_type: str = "",
        contract_number: str = "",
        hours: int = 16,
        education_form: str = "",
        education_format: str = "",
        custom_context: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None,
        template_name: Optional[str] = None
    ) -> str:
        """
        Generate an order document (enrollment, admission, or graduation).
        
        Args:
            order_type: Type of order ('enrollment', 'admission', 'graduation')
            listeners_data: List of listener dictionaries with keys:
                - full_name, position, court_name, region
            order_number: Order number (e.g., '111')
            order_date: Date of the order
            program_name: Name of the program
            stream_name: Name of the stream/group
            start_date: Start date of training (for enrollment)
            end_date: End date of training (for enrollment)
            contract_date: Contract date string (e.g., '28 мая 2025 года')
            contract_type: Contract type (e.g., 'государственным контрактом...')
            contract_number: Contract number (e.g., 'б/н')
            hours: Number of training hours
            education_form: Form of education (e.g., 'очной')
            education_format: Format of education (e.g., 'с применением...')
            custom_context: Additional template variables
            output_filename: Custom output filename
            template_name: Override template name (use file from templates/)
        
        Returns:
            Path to generated document
        """
        # Get template - either custom or from predefined types
        if template_name:
            template_path = self._resolve_template_path(template_name)
        else:
            if order_type not in self.TEMPLATE_TYPES:
                raise ValueError(f"Unknown order type: {order_type}. "
                               f"Available: {list(self.TEMPLATE_TYPES.keys())}")
            template_name = self.TEMPLATE_TYPES[order_type]
            template_path = self._resolve_template_path(template_name)
        
        if not template_path.exists():
            # Try generic order template as fallback
            generic_path = self._resolve_template_path('Приказ.docx')
            if generic_path.exists():
                template_path = generic_path
            else:
                # Generate basic order document without template
                return self._generate_basic_order(
                    order_type=order_type,
                    listeners_data=listeners_data,
                    order_number=order_number,
                    order_date=order_date,
                    program_name=program_name,
                    custom_context=custom_context,
                    output_filename=output_filename,
                )
        
        doc = DocxTemplate(template_path)
        
        # Prepare date parts
        order_parts = get_date_parts(order_date)
        
        # Prepare context
        context = {
            # Order details
            'order_number': order_number,
            'order_day': order_parts['day'],
            'order_month': order_parts['month'],
            'order_year': order_parts['year'],
            
            # Contract details
            'contract_type': contract_type,
            'contract_number': contract_number,
            'contract_date': contract_date,
            
            # Program details
            'program_name': program_name,
            'stream_name': stream_name,
            'hours': str(hours),
            'education_form': education_form,
            'education_form_genitive': self._get_education_form_genitive(education_form),
            'education_format': education_format,
        }

        # Sort listeners alphabetically by last_name (then full_name as fallback)
        sorted_listeners = sorted(
            listeners_data,
            key=lambda ld: (ld.get('last_name') or ld.get('full_name') or '').lower()
        )

        # Listeners list for table - add order_number and contract data
        enriched_listeners = []
        for idx, ld in enumerate(sorted_listeners, start=1):
            enriched = {**ld, 'order_number': idx}
            # Default empty contract fields
            enriched.setdefault('contract_number', '')
            enriched.setdefault('contract_date', '')
            enriched_listeners.append(enriched)

        # Enrich listeners with contract data from ContractJournal; also pull training_basis and grade_info
        program_id = (custom_context or {}).get('program_id')
        if program_id:
            try:
                with DatabaseSession() as session:
                    program_obj = session.query(Program).get(program_id)
                    if program_obj and program_obj.training_basis:
                        basis = program_obj.training_basis
                        basis_instr = self._decline_phrase(basis, 'instrumental')
                        context.setdefault('training_basis', basis)
                        context.setdefault('training_basis_instrumental', basis_instr)
                        context.setdefault(
                            'training_basis_phrase',
                            f"в соответствии с {basis_instr}" if basis_instr else ''
                        )
                    for el in enriched_listeners:
                        lid = el.get('id')
                        if lid:
                            contract = session.query(ContractJournal).filter(
                                ContractJournal.listener_id == lid,
                                ContractJournal.program_id == program_id,
                            ).order_by(ContractJournal.contract_number.desc()).first()
                            if contract:
                                el['contract_number'] = contract.contract_number
                                el['contract_date'] = (
                                    contract.contract_date.strftime('%d.%m.%Y')
                                    if contract.contract_date else ''
                                )
                            pl = session.query(ProgramListener).filter(
                                ProgramListener.listener_id == lid,
                                ProgramListener.program_id == program_id,
                            ).first()
                            el['grade_info'] = (pl.grade_info if pl else None) or ''
            except Exception:
                pass  # If lookup fails, leave defaults

        context['listeners'] = enriched_listeners
        context['listeners_count'] = len(listeners_data)
        
        # Add start/end dates if provided
        if start_date:
            context['start_date'] = format_date_russian(start_date)
        if end_date:
            context['end_date'] = format_date_russian(end_date)
        
        # Add service context
        context.update(self._prepare_service_context())
        
        # Add custom context
        if custom_context:
            context.update(custom_context)
        
        # Render document with jinja2 loop support
        import jinja2 as _jinja2
        jinja_env = _jinja2.Environment(extensions=['jinja2.ext.loopcontrols'])
        doc.render(context, jinja_env=jinja_env)
        
        # Generate output filename
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            type_names = {
                'enrollment': 'О_зачислении',
                'admission': 'О_допуске',
                'graduation': 'Об_отчислении',
                'thesis_topics': 'Об_утверждении_тем',
                'internship': 'О_направлении_на_стажировку',
                'theory_exam': 'О_допуске_к_экзамену',
                'theory_exam_retake': 'О_допуске_к_повторному_экзамену',
                'custom': 'Документ'
            }
            type_name = type_names.get(order_type, 'Документ')
            # Use template name for custom type
            if order_type == 'custom' and template_name:
                base_name = Path(template_name).stem
                output_name = f"{order_number}_{order_date.strftime('%d.%m.%Y')}_{base_name}.docx"
            else:
                output_name = f"{order_number}_{order_date.strftime('%d.%m.%Y')}_{type_name}.docx"
            output_path = self.output_dir / output_name
        
        doc.save(output_path)
        
        return str(output_path)

    def _generate_basic_order(
        self,
        order_type: str,
        listeners_data: List[Dict[str, Any]],
        order_number: str,
        order_date: datetime,
        program_name: str = "",
        custom_context: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None,
    ) -> str:
        """
        Generate a basic order Word document without a template.
        Uses python-docx directly to create the document.
        """
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        from ..services.order_journal_service import ORDER_TYPE_LABELS

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        order_type_label = ''
        if custom_context:
            order_type_label = custom_context.get('order_type_label', '')
        if not order_type_label:
            order_type_label = ORDER_TYPE_LABELS.get(order_type, 'Приказ')

        # Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run('ПРИКАЗ')
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        # Order details
        date_str = format_date_russian(order_date) if order_date else ''
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details_para.add_run(f'№ {order_number} от {date_str}')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        # Order type
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_para.add_run(f'«{order_type_label}»')
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        # Program
        if program_name:
            prog_para = doc.add_paragraph()
            prog_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = prog_para.add_run(f'по программе: {program_name}')
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'

        doc.add_paragraph()  # Empty line

        # Listeners table
        if listeners_data:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            headers = ['№ п/п', 'ФИО', 'Должность', 'Место работы']
            for i, header_text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header_text
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Data rows
            for idx, listener in enumerate(listeners_data, start=1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                full_name = listener.get('full_name', '')
                if not full_name:
                    parts = [
                        listener.get('last_name', ''),
                        listener.get('first_name', ''),
                        listener.get('middle_name', '') or '',
                    ]
                    full_name = ' '.join(p for p in parts if p)
                row.cells[1].text = full_name

                row.cells[2].text = listener.get('position', '') or ''
                row.cells[3].text = listener.get('workplace', '') or listener.get('court_name', '') or ''

                # Set font for all cells in the row
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'

            # Set column widths
            try:
                from docx.shared import Inches
                widths = [Cm(1.5), Cm(6), Cm(4), Cm(5)]
                for row in table.rows:
                    for idx_c, width in enumerate(widths):
                        row.cells[idx_c].width = width
            except Exception:
                pass

        # Generate output filename
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            type_names = {
                'enrollment': 'О_зачислении',
                'admission': 'О_допуске',
                'graduation': 'Об_отчислении',
                'thesis_topics': 'Об_утверждении_тем',
                'internship': 'О_направлении_на_стажировку',
                'theory_exam': 'О_допуске_к_экзамену',
                'theory_exam_retake': 'О_допуске_к_повторному_экзамену',
            }
            type_name = type_names.get(order_type, 'Приказ')
            output_name = f"{order_number}_{order_date.strftime('%d.%m.%Y')}_{type_name}.docx"
            output_path = self.output_dir / output_name

        self.output_dir.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return str(output_path)

    def generate_batch(
        self,
        listeners: List[Listener],
        program: Optional[Program],
        template_name: str,
        separate_files: bool = True,
        custom_context: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Generate documents for multiple listeners.
        
        Args:
            listeners: List of Listener instances
            program: Program instance (optional)
            template_name: Name of the template file
            separate_files: Generate separate file for each listener
            custom_context: Additional template variables
        
        Returns:
            List of paths to generated documents
        """
        generated_files = []
        
        for idx, listener in enumerate(listeners, start=1):
            try:
                file_path = self.generate_for_listener(
                    listener=listener,
                    program=program,
                    template_name=template_name,
                    order_number=idx,
                    custom_context=custom_context
                )
                generated_files.append(file_path)
            except Exception as e:
                # Log error but continue with other listeners
                print(f"Error generating document for {listener.full_name}: {e}")
        
        return generated_files
    
    def generate_with_table(
        self,
        listeners: List[Listener],
        program: Optional[Program],
        template_name: str,
        custom_context: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate a single document with a table of listeners.
        
        The template should contain a table with Jinja2 loop:
        {%tr for listener in listeners %}
        {{ listener.order_number }}
        {{ listener.full_name }}
        ...
        {%tr endfor %}
        
        Args:
            listeners: List of Listener instances
            program: Program instance (optional)
            template_name: Template filename
            custom_context: Additional template variables
            output_filename: Custom output filename
        
        Returns:
            Path to generated document
        """
        template_path = self._resolve_template_path(template_name)
        
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        
        doc = DocxTemplate(template_path)
        
        # Prepare listeners data for table
        listeners_data = []
        for idx, listener in enumerate(listeners, start=1):
            listener_context = self._prepare_listener_context(listener, idx)
            listeners_data.append(listener_context)
        
        # Prepare main context
        context = {
            'listeners': listeners_data,
            'listeners_count': len(listeners),
        }
        
        # Add program data
        if program:
            context.update(self._prepare_program_context(program))
        
        # Add service fields
        context.update(self._prepare_service_context())
        
        # Add custom context
        if custom_context:
            context.update(custom_context)
        
        # Render document
        doc.render(context)
        
        # Generate output filename
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            program_name = self._sanitize_filename(
                program.program_short_name or "program"
            ) if program else "document"
            output_name = f"{template_path.stem}_{program_name}_{timestamp}.docx"
            output_path = self.output_dir / output_name
        
        doc.save(output_path)
        
        return str(output_path)
    
    def _resolve_template_path(self, template_name: str) -> Path:
        """Resolve template path, adding extension if needed."""
        template_path = self.templates_dir / template_name
        
        if not template_path.suffix:
            template_path = template_path.with_suffix('.docx')
        
        return template_path
    
    def _prepare_context(
        self,
        listener: Listener,
        program: Optional[Program],
        order_number: int
    ) -> Dict[str, Any]:
        """
        Prepare complete template context.
        
        Args:
            listener: Listener instance
            program: Program instance
            order_number: Sequence number
        
        Returns:
            Dictionary with all template variables
        """
        context = {}
        
        # Add listener context with declensions
        context.update(self._prepare_listener_context(listener, order_number))
        
        # Add program context
        if program:
            context.update(self._prepare_program_context(program))
        
        # Add service context
        context.update(self._prepare_service_context())
        
        return context
    
    def _prepare_listener_context(
        self,
        listener: Listener,
        order_number: int
    ) -> Dict[str, Any]:
        """Prepare listener-specific context with all declensions."""
        # Get all name declensions
        context = self.declension.get_all_declensions(
            listener.last_name,
            listener.first_name,
            listener.middle_name
        )
        
        # Add other listener fields
        context['order_number'] = order_number
        context['position'] = listener.position or ''
        context['workplace'] = listener.workplace or ''
        context['court_name'] = listener.workplace or ''  # Alias for templates
        context['region'] = listener.region or ''
        context['notes'] = listener.notes or ''
        
        # Contact information
        context['mobile_phone'] = listener.mobile_phone or ''
        context['work_phone'] = listener.work_phone or ''
        context['email'] = listener.email or ''
        
        # Birth date
        if listener.birth_date:
            context['birth_date'] = listener.birth_date.strftime("%d.%m.%Y")
        else:
            context['birth_date'] = ''
        
        # Passport data
        context['passport_series_number'] = listener.passport_series_number or ''
        if listener.passport_issue_date:
            context['passport_issue_date'] = listener.passport_issue_date.strftime("%d.%m.%Y")
        else:
            context['passport_issue_date'] = ''
        context['passport_issued_by'] = listener.passport_issued_by or ''
        context['passport_department_code'] = listener.passport_department_code or ''
        
        # Addresses
        context['registration_address'] = listener.registration_address or ''
        context['actual_address'] = listener.actual_address or ''
        
        # Identifiers
        context['snils'] = listener.snils or ''
        context['inn'] = listener.inn or ''
        
        # Personal data consent
        context['personal_data_consent'] = 'Да' if listener.personal_data_consent else 'Нет'
        
        return context
    
    def _decline_phrase(self, phrase: str, case: str) -> str:
        """Decline each word of a phrase to the given case."""
        if not phrase:
            return ''
        words = phrase.split()
        return ' '.join(self.declension.decline_word(w, case) for w in words)

    def _prepare_program_context(self, program: Program) -> Dict[str, Any]:
        """Prepare program-specific context."""
        basis = program.training_basis or ''
        basis_instr = self._decline_phrase(basis, 'instrumental')
        return {
            'program_name': program.program_name or '',
            'program_short_name': program.program_short_name or '',
            'training_basis': basis,
            'training_basis_instrumental': basis_instr,
            'training_basis_phrase': f"в соответствии с {basis_instr}" if basis_instr else '',
            'training_period': program.training_period or '',
            'program_volume': program.program_volume or '',
            'education_form': program.education_form or '',
            'education_format': program.education_format or '',
            'listener_category': program.listener_category or '',
            'expulsion_date': program.formatted_expulsion_date,
        }
    
    def _prepare_service_context(self) -> Dict[str, Any]:
        """Prepare service context (dates, etc.)."""
        now = datetime.now()
        
        return {
            'current_date': now.strftime("%d.%m.%Y"),
            'current_year': str(now.year),
            'current_month': now.strftime("%m"),
            'current_day': now.strftime("%d"),
            'generation_datetime': now.strftime("%d.%m.%Y %H:%M:%S"),
        }

    @staticmethod
    def _get_education_form_genitive(form: str) -> str:
        """Get genitive case of education form (родительный падеж формы обучения)."""
        mapping = {
            'очная': 'очной',
            'заочная': 'заочной',
            'очно-заочная': 'очно-заочной',
            'Очная': 'Очной',
            'Заочная': 'Заочной',
            'Очно-заочная': 'Очно-заочной',
        }
        return mapping.get(form, form)

    def _sanitize_filename(self, name: str) -> str:
        """Remove invalid characters from filename."""
        if not name:
            return "document"
        
        # Replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        result = name
        for char in invalid_chars:
            result = result.replace(char, '_')
        
        # Remove leading/trailing spaces and dots
        result = result.strip(' .')
        
        return result[:50] if result else "document"
    
    def _determine_document_type(self, template_name: str) -> str:
        """
        Determine document type from template name for registration.
        
        Args:
            template_name: Name of the template file
        
        Returns:
            Document type string (enrollment, admission, graduation, order, other)
        """
        name_lower = template_name.lower()
        
        # Check for enrollment (зачисление)
        if any(word in name_lower for word in ['зачисл', 'enrollment', 'enroll']):
            return 'enrollment'
        
        # Check for admission (допуск)
        if any(word in name_lower for word in ['допуск', 'admission', 'admit']):
            return 'admission'
        
        # Check for graduation/expulsion (отчисление)
        if any(word in name_lower for word in ['отчисл', 'graduation', 'expuls', 'выпуск']):
            return 'graduation'
        
        # Check for order (приказ)
        if any(word in name_lower for word in ['приказ', 'order', 'распоряж']):
            return 'order'
        
        return 'other'
    
    def get_template_variables(self, template_name: str) -> List[str]:
        """
        Extract variable names from a template.
        Note: This is a basic extraction, may not catch all variables.
        
        Args:
            template_name: Template filename
        
        Returns:
            List of variable names found in template
        """
        import re
        
        template_path = self._resolve_template_path(template_name)
        
        if not template_path.exists():
            return []
        
        try:
            doc = DocxTemplate(template_path)
            
            # Get the document XML
            xml_content = doc.get_docx().element.body.xml
            
            # Find all Jinja2 variables
            pattern = r'\{\{[\s]*([a-zA-Z_][a-zA-Z0-9_]*)[\s]*\}\}'
            variables = re.findall(pattern, xml_content)
            
            return sorted(set(variables))
        except Exception:
            return []


    def fix_template_loop(self, template_name: str) -> dict:
        """
        Add {% for listener in listeners %} loop to a template if missing.
        
        Args:
            template_name: Template filename
        
        Returns:
            dict with keys: 'status' ('fixed', 'already_ok', 'no_vars', 'error'),
                            'message' (str)
        """
        import re as _re
        import zipfile as _zipfile
        import shutil as _shutil

        tpl_path = self._resolve_template_path(template_name)

        if not tpl_path.exists():
            return {'status': 'error', 'message': f'Файл не найден: {tpl_path}'}

        try:
            with _zipfile.ZipFile(tpl_path, 'r') as z:
                xml = z.read('word/document.xml').decode('utf-8')

            # Already has the loop
            if '{% for listener' in xml or '{%- for listener' in xml:
                return {'status': 'already_ok', 'message': 'Цикл уже есть'}

            # Find table row containing listener/loop variables
            tr_pattern = r'(<w:tr\b[^>]*>.*?</w:tr>)'
            matches = list(_re.finditer(tr_pattern, xml, _re.DOTALL))

            found = False
            for match in matches:
                tr_content = match.group(1)
                if 'loop' in tr_content or 'listener' in tr_content:
                    new_tr = '{% for listener in listeners %}' + tr_content + '{% endfor %}'
                    xml = xml.replace(tr_content, new_tr, 1)
                    found = True
                    break

            if not found:
                return {
                    'status': 'no_vars',
                    'message': 'Нет строки таблицы с переменными listener/loop'
                }

            # Save modified template
            tmp_path = str(tpl_path) + '.tmp'
            with _zipfile.ZipFile(tpl_path, 'r') as zin:
                with _zipfile.ZipFile(tmp_path, 'w', _zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        if item.filename == 'word/document.xml':
                            zout.writestr(item, xml.encode('utf-8'))
                        else:
                            zout.writestr(item, zin.read(item.filename))

            _shutil.move(tmp_path, str(tpl_path))

            return {'status': 'fixed', 'message': 'Цикл добавлен'}

        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    def fix_all_templates(self) -> List[dict]:
        """
        Fix all templates in templates/ folder by adding the
        {% for listener in listeners %} loop where missing.
        
        Returns:
            List of dicts: [{'name': str, 'status': str, 'message': str}, ...]
        """
        templates = self.get_available_templates()
        results = []
        for tpl in templates:
            res = self.fix_template_loop(tpl)
            results.append({'name': tpl, **res})
        return results


class DocumentGeneratorError(Exception):
    """Exception for document generation errors."""
    pass


# Helper functions for preparing Russian dates
MONTHS_GENITIVE = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
}


def format_date_russian(date: datetime) -> str:
    """Format date in Russian style: '12 ноября 2025 г.'"""
    return f"{date.day} {MONTHS_GENITIVE[date.month]} {date.year} г."


def get_date_parts(date: datetime) -> Dict[str, str]:
    """Get date parts for templates."""
    return {
        'day': str(date.day),
        'month': MONTHS_GENITIVE[date.month],
        'year': str(date.year),
    }
