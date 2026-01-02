#!/usr/bin/env python3
"""
Script to create Word templates with Jinja2 placeholders for docxtpl.
Creates three templates: enrollment, admission, graduation.
"""

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_width(cell, width_cm):
    """Set cell width in cm."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 567 twips per cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def create_enrollment_template():
    """Create 'О зачислении' (Enrollment) template."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Header
    for text, bold in [
        ('ВОСТОЧНО-СИБИРСКИЙ ФИЛИАЛ', True),
        ('федерального государственного бюджетного образовательного учреждения высшего образования', False),
        ('«РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ПРАВОСУДИЯ', True),
        ('ИМЕНИ В.М. ЛЕБЕДЕВА»', True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
    
    doc.add_paragraph()
    
    # ПРИКАЗ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПРИКАЗ')
    run.bold = True
    
    # Date and number
    p = doc.add_paragraph()
    p.add_run('«{{ order_day }}» {{ order_month }} {{ order_year }} г.')
    p.add_run('\t' * 7)
    p.add_run('№ {{ order_number }}')
    
    p = doc.add_paragraph()
    p.add_run('г. Иркутск')
    
    doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('О зачислении')
    run.bold = True
    
    doc.add_paragraph()
    
    # Main text
    p = doc.add_paragraph()
    p.add_run('В соответствии с государственным контрактом на оказание услуг по повышению квалификации б/н от {{ contract_date }}  ')
    run = p.add_run('п р и к а з ы в а ю :')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('1. Зачислить в ВСФ ФГБОУВО «РГУП им. В.М. Лебедева» на период с {{ start_date }} по {{ end_date }} в состав слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов) потока {{ stream_name }} по дополнительной профессиональной программе повышения квалификации «{{ program_name }}» лиц согласно приложению № 1.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    p = doc.add_paragraph()
    p.add_run('Первый заместитель директора')
    p.add_run('\t' * 5)
    p.add_run('Е.Ю. Рузавина')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Проект вносит:')
    p = doc.add_paragraph()
    p.add_run('Заместитель декана факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов)')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Т.В. Жиленкова')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Исполнитель:')
    p = doc.add_paragraph()
    p.add_run('Специалист по учебной работе 1 категории факультета повышения квалификации и профессиональной переподготовки судей, государственных гражданских служащих судов и Судебного департамента')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Л.Н. Комлева')
    
    # Page break for Appendix
    doc.add_page_break()
    
    # Appendix header
    for text in [
        'Приложение № 1',
        'к приказу ВСФ ФГБОУВО «РГУП',
        'им. В.М. Лебедева»',
        'от {{ order_day }} {{ order_month }} {{ order_year }} г. № {{ order_number }}'
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(text)
    
    doc.add_paragraph()
    
    # List title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Список слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов), зачисленных на поток {{ stream_name }} по дополнительной профессиональной программе повышения квалификации «{{ program_name }}»')
    
    doc.add_paragraph()
    
    # Table with docxtpl loop syntax
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['№ п/п', 'Фамилия Имя Отчество', 'Должность', 'Наименование суда', 'Наименование субъекта Российской Федерации']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    
    # Data row with docxtpl loop - using special {%tr syntax
    row = table.add_row()
    row.cells[0].text = '{%tr for listener in listeners %}{{ loop.index }}'
    row.cells[1].text = '{{ listener.full_name }}'
    row.cells[2].text = '{{ listener.position }}'
    row.cells[3].text = '{{ listener.court_name }}'
    row.cells[4].text = '{{ listener.region }}{%tr endfor %}'
    
    return doc


def create_admission_template():
    """Create 'О допуске к аттестации' (Admission) template."""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Header
    for text, bold in [
        ('ВОСТОЧНО-СИБИРСКИЙ ФИЛИАЛ', True),
        ('федерального государственного бюджетного образовательного учреждения высшего образования', False),
        ('«РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ПРАВОСУДИЯ', True),
        ('ИМЕНИ В.М. ЛЕБЕДЕВА»', True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПРИКАЗ')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('«{{ order_day }}» {{ order_month }} {{ order_year }} г.')
    p.add_run('\t' * 7)
    p.add_run('№ {{ order_number }}')
    
    p = doc.add_paragraph()
    p.add_run('г. Иркутск')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('О допуске слушателей')
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('к итоговой аттестации')
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('В связи с успешным выполнением учебного плана обучения по дополнительной профессиональной программе ')
    run = p.add_run('п р и к а з ы в а ю :')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('1. Допустить слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов) потока {{ stream_name }} по дополнительной профессиональной программе повышения квалификации «{{ program_name }}» к итоговой аттестации согласно приложению № 1.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    p = doc.add_paragraph()
    p.add_run('Первый заместитель директора')
    p.add_run('\t' * 5)
    p.add_run('Е.Ю. Рузавина')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Проект вносит:')
    p = doc.add_paragraph()
    p.add_run('Заместитель декана факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов)')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Т.В. Жиленкова')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Исполнитель:')
    p = doc.add_paragraph()
    p.add_run('Специалист по учебной работе 1 категории факультета повышения квалификации и профессиональной переподготовки судей, государственных гражданских служащих судов и Судебного департамента')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Л.Н. Комлева')
    
    doc.add_page_break()
    
    # Appendix
    for text in [
        'Приложение № 1',
        'к приказу ВСФ ФГБОУВО «РГУП',
        'им. В.М. Лебедева»',
        'от {{ order_day }} {{ order_month }} {{ order_year }} г. № {{ order_number }}'
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(text)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Список слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов), допущенных к итоговой аттестации по дополнительной профессиональной программе повышения квалификации «{{ program_name }}»')
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['№ п/п', 'Фамилия Имя Отчество', 'Должность', 'Наименование суда', 'Наименование субъекта Российской Федерации']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    
    row = table.add_row()
    row.cells[0].text = '{%tr for listener in listeners %}{{ loop.index }}'
    row.cells[1].text = '{{ listener.full_name }}'
    row.cells[2].text = '{{ listener.position }}'
    row.cells[3].text = '{{ listener.court_name }}'
    row.cells[4].text = '{{ listener.region }}{%tr endfor %}'
    
    return doc


def create_graduation_template():
    """Create 'Об отчислении' (Graduation) template."""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Header
    for text, bold in [
        ('ВОСТОЧНО-СИБИРСКИЙ ФИЛИАЛ', True),
        ('федерального государственного бюджетного образовательного учреждения высшего образования', False),
        ('«РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ПРАВОСУДИЯ', True),
        ('ИМЕНИ В.М. ЛЕБЕДЕВА»', True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПРИКАЗ')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('«{{ order_day }}» {{ order_month }} {{ order_year }} г.')
    p.add_run('\t' * 7)
    p.add_run('№ {{ order_number }}')
    
    p = doc.add_paragraph()
    p.add_run('г. Иркутск')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Об отчислении слушателей')
    run.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('В соответствии с государственным контрактом на оказание услуг по повышению квалификации б/н от {{ contract_date }}  ')
    run = p.add_run('п р и к а з ы в а ю:')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('1. Отчислить слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов) потока {{ stream_name }} по дополнительной профессиональной программе повышения квалификации «{{ program_name }}» в связи с завершением обучения согласно приложению № 1.')
    
    p = doc.add_paragraph()
    p.add_run('2. Выдать слушателям, успешно выполнившим учебный план, удостоверения о повышении квалификации в объеме {{ hours }} часов.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    p = doc.add_paragraph()
    p.add_run('Первый заместитель директора')
    p.add_run('\t' * 5)
    p.add_run('Е.Ю. Рузавина')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Проект вносит:')
    p = doc.add_paragraph()
    p.add_run('Заместитель декана факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов)')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Т.В. Жиленкова')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Исполнитель:')
    p = doc.add_paragraph()
    p.add_run('Специалист по учебной работе 1 категории факультета повышения квалификации и профессиональной переподготовки судей, государственных гражданских служащих судов и Судебного департамента')
    p = doc.add_paragraph()
    p.add_run('+7(3952)41-07-40')
    p.add_run('\t' * 6)
    p.add_run('Л.Н. Комлева')
    
    doc.add_page_break()
    
    # Appendix
    for text in [
        'Приложение № 1',
        'к приказу ВСФ ФГБОУВО «РГУП',
        'им. В.М. Лебедева»',
        'от {{ order_day }} {{ order_month }} {{ order_year }} г. № {{ order_number }}'
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(text)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Список слушателей факультета повышения квалификации и переподготовки судей, государственных гражданских служащих судов и Судебного департамента (ФПК судей и госслужащих судов), завершивших обучение по дополнительной профессиональной программе повышения квалификации «{{ program_name }}»')
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['№ п/п', 'Фамилия Имя Отчество', 'Должность', 'Наименование суда', 'Наименование субъекта Российской Федерации']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    
    row = table.add_row()
    row.cells[0].text = '{%tr for listener in listeners %}{{ loop.index }}'
    row.cells[1].text = '{{ listener.full_name }}'
    row.cells[2].text = '{{ listener.position }}'
    row.cells[3].text = '{{ listener.court_name }}'
    row.cells[4].text = '{{ listener.region }}{%tr endfor %}'
    
    return doc


def main():
    """Create all templates."""
    os.makedirs('templates', exist_ok=True)
    
    # Create templates
    templates = [
        ('enrollment_template.docx', create_enrollment_template, 'О зачислении'),
        ('admission_template.docx', create_admission_template, 'О допуске к аттестации'),
        ('graduation_template.docx', create_graduation_template, 'Об отчислении'),
    ]
    
    for filename, creator, name in templates:
        doc = creator()
        path = f'templates/{filename}'
        doc.save(path)
        print(f'✓ Создан шаблон "{name}": {path}')
    
    print()
    print('Все шаблоны созданы!')
    print()
    print('Доступные плейсхолдеры:')
    print('  {{ order_day }}      - день приказа (12)')
    print('  {{ order_month }}    - месяц приказа (ноября)')
    print('  {{ order_year }}     - год приказа (2025)')
    print('  {{ order_number }}   - номер приказа (111)')
    print('  {{ contract_date }}  - дата контракта (28 мая 2025 года)')
    print('  {{ start_date }}     - дата начала (12 ноября 2025 г.)')
    print('  {{ end_date }}       - дата окончания (18 ноября 2025 г.)')
    print('  {{ stream_name }}    - название потока')
    print('  {{ program_name }}   - название программы')
    print('  {{ hours }}          - количество часов')
    print()
    print('В таблице (цикл по listeners):')
    print('  {{ listener.full_name }}   - ФИО')
    print('  {{ listener.position }}    - должность')
    print('  {{ listener.court_name }}  - наименование суда')
    print('  {{ listener.region }}      - регион')


if __name__ == '__main__':
    main()
